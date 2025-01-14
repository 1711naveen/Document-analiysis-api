import re
from docx.shared import RGBColor
from num2words import num2words
from word2number import w2n
import enchant
from fastapi import APIRouter, HTTPException, Depends, Query
from fastapi.responses import FileResponse
import os
import docx
from sqlalchemy import text
from db_config import get_db_connection
import mammoth
from datetime import datetime
from pathlib import Path
import logging  
import roman
from roman import fromRoman
from urllib.parse import urlparse

router = APIRouter()

# us_dict = enchant.Dict("en_US")

us_dict = enchant.DictWithPWL("en_US","mywords.txt")

global_logs = []

def fetch_abbreviation_mappings():
    conn = get_db_connection()
    cursor = conn.cursor()
    cursor.execute("SELECT original_word, abbreviated_form FROM abbreviation_mapping")
    mappings = cursor.fetchall()
    conn.close()
    return {row[0]: row[1] for row in mappings}


century_map = {
    1: "first",
    2: "second",
    3: "third",
    4: "fourth",
    5: "fifth",
    6: "sixth",
    7: "seventh",
    8: "eighth",
    9: "ninth",
    10: "tenth",
    11: "eleventh",
    12: "twelfth",
    13: "thirteenth",
    14: "fourteenth",
    15: "fifteenth",
    16: "sixteenth",
    17: "seventeenth",
    18: "eighteenth",
    19: "nineteenth",
    20: "twentieth",
    21: "twenty-first",
    22: "twenty-second",
    23: "twenty-third",
    24: "twenty-fourth",
    25: "twenty-fifth",
}

def apply_abbreviation_mapping(text, abbreviation_dict, line_number):
    global global_logs
    words = text.split()
    updated_text = []
    for word in words:
        updated_word = abbreviation_dict.get(word, word)
        if word != updated_word:
            global_logs.append(f"[apply_abbreviation_mapping] Line {line_number}: '{word}' -> '{updated_word}'")
        updated_text.append(updated_word)
    return ' '.join(updated_text)



def apply_number_abbreviation_rule(text, line_number):
    """
    Replaces 'Number X' or 'number X' with 'No. X' or 'no. X' and logs changes.
    :param text: The input text.
    :param line_number: Line number for logging.
    :return: Updated text with number abbreviations applied.
    """
    global global_logs

    def replace_number(match):
        word = match.group(1)
        num = match.group(2)
        updated_text = f"No. {num}" if word.istitle() else f"no. {num}"
        global_logs.append(f"[apply_number_abbreviation_rule] Line {line_number}: '{match.group(0)}' -> '{updated_text}'")
        return updated_text

    pattern = r'\b(Number|number)\s(\d+)\b'
    return re.sub(pattern, replace_number, text)



# Done
# def convert_century(text, line_number_offset):
#     """
#     Converts century notation like '21st' to 'the twenty-first century'
#     and logs the changes with line numbers.
    
#     :param text: The entire text to process, possibly spanning multiple lines.
#     :param line_number_offset: The starting line number for this chunk of text.
#     :return: The updated text with century notations converted.
#     """
#     global global_logs  # Global log to record changes
#     lines = text.split('\n')  # Split text into individual lines
#     updated_lines = []

#     for index, line in enumerate(lines):
#         words = line.split()  # Split line into words
#         for i, word in enumerate(words):
#             match = re.match(r"(\d+)(st|nd|rd|th)$", word)  # Match century notation
#             if match:
#                 num = int(match.group(1))
#                 if num in century_map:
#                     # Original and converted word
#                     original_word = match.group(0)
#                     converted_word = f"the {century_map[num]} century"
                    
#                     # Log the change with the actual line number
#                     global_logs.append(
#                         f"[convert century] Line {line_number_offset + index}: {original_word} -> {converted_word}"
#                     )
                    
#                     # Replace the word in the line
#                     words[i] = converted_word
        
#         # Rebuild the updated line
#         updated_lines.append(' '.join(words))

#     # Return the updated text with all lines rebuilt
#     return '\n'.join(updated_lines)



def convert_century(paragraph, line_number_offset):
    """
    Converts century notation like '21st' to 'the twenty-first century'
    in a paragraph's runs, preserving formatting, and logs the changes with line numbers.

    :param paragraph: The paragraph object containing runs.
    :param line_number_offset: The starting line number for this paragraph.
    """
    global global_logs  # Global log to record changes
    century_map = {
        1: "first", 2: "second", 3: "third", 4: "fourth", 5: "fifth",
        6: "sixth", 7: "seventh", 8: "eighth", 9: "ninth", 10: "tenth",
        11: "eleventh", 12: "twelfth", 13: "thirteenth", 14: "fourteenth",
        15: "fifteenth", 16: "sixteenth", 17: "seventeenth", 18: "eighteenth",
        19: "nineteenth", 20: "twentieth", 21: "twenty-first",
        22: "twenty-second", 23: "twenty-third", 24: "twenty-fourth",
        25: "twenty-fifth"  # Add more as needed
    }

    for run in paragraph.runs:
        words = run.text.split()
        for i, word in enumerate(words):
            match = re.match(r"(\d+)(st|nd|rd|th)$", word)  # Match century notation
            if match:
                num = int(match.group(1))
                if num in century_map:
                    original_word = match.group(0)
                    converted_word = f"the {century_map[num]} century"
                    
                    # Log the change with the actual line number
                    global_logs.append(
                        f"[convert century] Line {line_number_offset}: {original_word} -> {converted_word}"
                    )
                    
                    # Replace the word in the run
                    words[i] = converted_word
        
        # Update the run's text while preserving formatting
        run.text = ' '.join(words)



def clean_word(word):
    return word.strip(",.?!:;\"'()[]{}")



# Done
def replace_curly_quotes_with_straight(text):
    return text.replace("“", '"').replace("”", '"').replace("‘", "'").replace("’", "'")

    
    

# def replace_straight_quotes_with_curly(text):
#     # Replace straight double quotes with opening and closing curly quotes
#     text = re.sub(r'(^|[\s([{])"', r'\1“', text)  # Opening double quotes
#     text = re.sub(r'"', r'”', text)  # Closing double quotes
    
#     # Replace straight single quotes with opening and closing curly quotes
#     text = re.sub(r"(^|[\s([{])'", r'\1‘', text)  # Opening single quotes
#     text = re.sub(r"'", r'’', text)  # Closing single quotes
    
#     text = re.sub(r"([a-zA-Z]+)'([a-zA-Z]+)", r"\1‘\2", text)  # Curly starting single quote after word
    
#     return text


def replace_straight_quotes_with_curly(paragraph):
    """
    Replaces straight quotes with curly quotes in a paragraph's runs while preserving formatting.
    
    :param paragraph: The paragraph object containing runs.
    """
    for run in paragraph.runs:
        text = run.text

        # Replace straight double quotes with opening and closing curly quotes
        text = re.sub(r'(^|[\s([{])"', r'\1“', text)  # Opening double quotes
        text = re.sub(r'"', r'”', text)  # Closing double quotes

        # Replace straight single quotes with opening and closing curly quotes
        text = re.sub(r"(^|[\s([{])'", r'\1‘', text)  # Opening single quotes
        text = re.sub(r"'", r'’', text)  # Closing single quotes

        # Handle curly starting single quotes within words
        text = re.sub(r"([a-zA-Z]+)'([a-zA-Z]+)", r"\1‘\2", text)

        # Update the run's text while preserving formatting
        run.text = text





# Done
# def correct_acronyms(text, line_number):
#     global global_logs
#     original_text = text
#     words = text.split()
#     corrected_words = []
#     for word in words:
#         original_word = word
#         if re.match(r"([a-z]\.){2,}[a-z]\.?", word):
#             word = word.replace(".", "")
#         elif re.match(r"([A-Z]\.){2,}[A-Z]\.?", word):
#             word = word.replace(".", "")
#         if word != original_word:
#             global_logs.append(
#                 f"[correct_acronyms] Line {line_number}: '{original_word}' -> '{word}'"
#             )
#         corrected_words.append(word)
#     corrected_text = " ".join(corrected_words)
#     return corrected_text

def correct_acronyms(paragraph, line_number):
    """
    Corrects acronyms in a paragraph's runs, preserving formatting, and logs changes with line numbers.
    
    :param paragraph: The paragraph object containing runs.
    :param line_number: Line number for logging.
    """
    global global_logs
    for run in paragraph.runs:
        original_text = run.text
        words = original_text.split()
        corrected_words = []
        
        for word in words:
            original_word = word
            if re.match(r"([a-z]\.){2,}[a-z]\.?", word):  # Matches lowercase acronyms
                word = word.replace(".", "")
            elif re.match(r"([A-Z]\.){2,}[A-Z]\.?", word):  # Matches uppercase acronyms
                word = word.replace(".", "")
            
            if word != original_word:
                global_logs.append(
                    f"[correct_acronyms] Line {line_number}: '{original_word}' -> '{word}'"
                )
            
            corrected_words.append(word)
        
        # Update the run's text while preserving formatting
        run.text = " ".join(corrected_words)



# def enforce_am_pm(text, line_num):
#     """
#     Ensures consistent formatting for 'am' and 'pm' in the entire paragraph and logs changes.
#     :param text: The paragraph text to process.
#     :param line_num: The line number in the document for logging.
#     :return: The updated text with corrected 'am' and 'pm' formats.
#     """
#     global global_logs  # Use a global log to record changes
#     original_text = text  # Store the original text for comparison
#     words = text.split()  # Split the paragraph into words

#     corrected_words = []
#     for word in words:
#         original_word = word  # Store the original word for logging
#         word_lower = word.lower()  # Convert word to lowercase for comparison

#         # Check and correct 'am' or 'pm' formats
#         if word_lower in {"am", "a.m", "pm", "p.m"}:
#             if "a" in word_lower:
#                 corrected_word = "a.m."
#             elif "p" in word_lower:
#                 corrected_word = "p.m."
            
#             # Log the change if the word was modified
#             if corrected_word != original_word:
#                 global_logs.append(
#                     f"[am pm change] Line {line_num}: '{original_word}' -> '{corrected_word}'"
#                 )
#         else:
#             corrected_word = word  # Keep the word unchanged if no match

#         corrected_words.append(corrected_word)  # Add the corrected word to the list

#     # Join the corrected words to form the updated paragraph
#     corrected_text = " ".join(corrected_words)

#     return corrected_text


def enforce_am_pm(paragraph, line_num):
    """
    Ensures consistent formatting for 'am' and 'pm' in the entire paragraph's runs, preserving formatting, and logs changes.
    
    :param paragraph: The paragraph object containing runs.
    :param line_num: The line number in the document for logging.
    """
    global global_logs  # Use a global log to record changes
    
    for run in paragraph.runs:
        original_text = run.text  # Store the original text of the run
        words = original_text.split()  # Split the run text into words

        corrected_words = []
        for word in words:
            original_word = word  # Store the original word for logging
            word_lower = word.lower()  # Convert word to lowercase for comparison

            # Check and correct 'am' or 'pm' formats
            if word_lower in {"am", "a.m", "pm", "p.m"}:
                if "a" in word_lower:
                    corrected_word = "a.m."
                elif "p" in word_lower:
                    corrected_word = "p.m."
                
                # Log the change if the word was modified
                if corrected_word != original_word:
                    global_logs.append(
                        f"[am pm change] Line {line_num}: '{original_word}' -> '{corrected_word}'"
                    )
            else:
                corrected_word = word  # Keep the word unchanged if no match

            corrected_words.append(corrected_word)  # Add the corrected word to the list

        # Update the run's text while preserving formatting
        run.text = " ".join(corrected_words)



# Done
# [apostrophes change] : 60's -> 1960s 
# def remove_unnecessary_apostrophes(word, line_num):
#     original_word = word
#     global global_logs
#     word = re.sub(r"(\d{4})'s\b", r"\1s", word)
#     word = re.sub(r"'(\d{2})s\b", r"\1s", word)
#     word = re.sub(r"(\d{4}s)'\b", r"\1", word)
#     word = re.sub(r"(\d+)'(s|st|nd|rd|th)\b", r"\1\2", word)
#     word = re.sub(r"^(\d{2})s\b", r"19\1s", word)
#     if word != original_word:
#         global_logs.append(f"[apostrophes change] Line {line_num}: {original_word} -> {word}")
    
#     return word


def remove_unnecessary_apostrophes(paragraph, line_num):
    """
    Removes unnecessary apostrophes in a paragraph's runs, preserving formatting, and logs changes.
    
    :param paragraph: The paragraph object containing runs.
    :param line_num: The line number in the document for logging.
    """
    global global_logs  # Use a global log to record changes
    
    for run in paragraph.runs:
        original_text = run.text  # Store the original text of the run
        words = original_text.split()  # Split the run text into words
        corrected_words = []

        for word in words:
            original_word = word  # Store the original word for logging

            # Apply regex transformations to remove unnecessary apostrophes
            word = re.sub(r"(\d{4})'s\b", r"\1s", word)  # Handle '1980's' -> '1980s'
            word = re.sub(r"'(\d{2})s\b", r"\1s", word)  # Handle '70's' -> '70s'
            word = re.sub(r"(\d{4}s)'\b", r"\1", word)  # Handle '1980s'' -> '1980s'
            word = re.sub(r"(\d+)'(s|st|nd|rd|th)\b", r"\1\2", word)  # Handle '1980'th -> '1980th'
            word = re.sub(r"^(\d{2})s\b", r"19\1s", word)  # Handle '80s' -> '1980s'

            # Log the change if the word was modified
            if word != original_word:
                global_logs.append(
                    f"[apostrophes change] Line {line_num}: {original_word} -> {word}"
                )

            corrected_words.append(word)

        # Update the run's text while preserving formatting
        run.text = " ".join(corrected_words)





# Not working
def spell_out_number_and_unit_with_rules(sentence, line_number):
    global global_logs
    original_words = sentence.split()
    modified_words = original_words[:]
    unit_pattern = r"(\d+)\s+([a-zA-Z]+)"
    number_pattern = r"\b(\d+)\b"

    for i, word in enumerate(original_words):
        # Handle number followed by unit
        if re.match(unit_pattern, " ".join(original_words[i:i+2])):
            continue  # Skip since it's already formatted correctly
        # Spell out numbers less than 10
        elif re.match(number_pattern, word):
            number = int(word)
            if number < 10:
                modified_words[i] = num2words(number, to="cardinal")
    
    # Log only changes
    for orig, mod in zip(original_words, modified_words):
        if orig != mod:
            global_logs.append(f"[spell_out_number_and_unit_with_rules] Line {line_number}: '{orig}' -> '{mod}'")
    return " ".join(modified_words)



# def use_numerals_with_percent(text):
#     global global_logs

#     lines = text.splitlines()
#     modified_text = []

#     for line_number, line in enumerate(lines, 1):
#         original_line = line
#         modified_line = line
#         def replace_spelled_out_percent(match):
#             word = match.group(1)
#             try:
#                 num = w2n.word_to_num(word.lower())
#                 modified = f"{num}%"
#                 global_logs.append(
#                     f"[numerals with percent] Line {line_number}: '{word} percent' -> '{modified}'"
#                 )
#                 return modified
#             except ValueError:
#                 return match.group(0)

#         modified_line = re.sub(
#             r"\b([a-zA-Z\s\-]+)\s?(percent|per cent|percentage)\b",
#             replace_spelled_out_percent,
#             modified_line,
#             flags=re.IGNORECASE,
#         )

#         def replace_numerical_percent(match):
#             number = match.group(1)
#             modified = f"{number}%"
#             global_logs.append(
#                 f"[numerals with percent] Line {line_number}: '{match.group(0)}' -> '{modified}'"
#             )
#             return modified

#         modified_line = re.sub(
#             r"(\d+)\s?(percent|per cent|percentage)\b", replace_numerical_percent, modified_line, flags=re.IGNORECASE
#         )

#         modified_text.append(modified_line)

#     return "\n".join(modified_text)


def use_numerals_with_percent(paragraph, line_number):
    """
    Ensures percentages are represented as numerals followed by '%' in a paragraph's runs, preserving formatting.
    Logs changes.
    
    :param paragraph: The paragraph object containing runs.
    :param line_number: The line number in the document for logging.
    """
    global global_logs  # Use a global log to record changes

    for run in paragraph.runs:
        original_text = run.text  # Store the original text of the run
        corrected_text = original_text

        # Replace spelled-out percentages with numeral equivalents
        def replace_spelled_out_percent(match):
            word = match.group(1)
            try:
                num = w2n.word_to_num(word.lower())  # Convert word to number
                modified = f"{num}%"
                global_logs.append(
                    f"[numerals with percent] Line {line_number}: '{word} percent' -> '{modified}'"
                )
                return modified
            except ValueError:
                return match.group(0)

        corrected_text = re.sub(
            r"\b([a-zA-Z\s\-]+)\s?(percent|per cent|percentage)\b",
            replace_spelled_out_percent,
            corrected_text,
            flags=re.IGNORECASE,
        )

        # Replace numerical percentages with a consistent '%'
        def replace_numerical_percent(match):
            number = match.group(1)
            modified = f"{number}%"
            global_logs.append(
                f"[numerals with percent] Line {line_number}: '{match.group(0)}' -> '{modified}'"
            )
            return modified

        corrected_text = re.sub(
            r"(\d+)\s?(percent|per cent|percentage)\b",
            replace_numerical_percent,
            corrected_text,
            flags=re.IGNORECASE,
        )

        # Update the run's text while preserving formatting
        run.text = corrected_text




# def enforce_eg_rule_with_logging(text):
#     lines = text.splitlines()
#     updated_lines = []
#     for line_number, line in enumerate(lines, start=1):
#         original_line = line

#         # Step 1: Match "eg" or "e.g." with optional surrounding spaces and punctuation
#         new_line = re.sub(r'\beg\b', 'e.g.', line, flags=re.IGNORECASE)
#         new_line = re.sub(r'\beg,\b', 'e.g.', new_line, flags=re.IGNORECASE)  # Handle "eg,"

#         # Step 2: Fix extra periods like `e.g..` or `e.g...,` and ensure proper punctuation
#         new_line = re.sub(r'\.([.,])', r'\1', new_line)  # Removes an extra period before a comma or period
#         new_line = re.sub(r'\.\.+', '.', new_line)  # Ensures only one period after e.g.

#         # Step 3: Remove comma if e.g... is followed by it (e.g..., -> e.g.)
#         new_line = re.sub(r'e\.g\.,', 'e.g.', new_line)

#         # Step 4: Change e.g, to e.g.
#         new_line = re.sub(r'e\.g,', 'e.g.', new_line)

#         # Log changes if the line is updated
#         if new_line != line:
#             global_logs.append(
#                 f"[e.g. correction] Line {line_number}: {line.strip()} -> {new_line.strip()}"
#             )
        
#         updated_lines.append(new_line)
#     return "\n".join(updated_lines)


def enforce_eg_rule_with_logging(paragraph, line_number):
    """
    Ensures consistent usage of 'e.g.' in the paragraph's runs and logs changes,
    preserving the formatting.
    
    :param paragraph: The paragraph object containing runs.
    :param line_number: The line number in the document for logging.
    """
    global global_logs  # Use a global log to record changes

    for run in paragraph.runs:
        original_text = run.text  # Store the original text of the run
        corrected_text = original_text

        # Step 1: Match "eg" or "e.g." with optional surrounding spaces and punctuation
        corrected_text = re.sub(r'\beg\b', 'e.g.', corrected_text, flags=re.IGNORECASE)
        corrected_text = re.sub(r'\beg,\b', 'e.g.', corrected_text, flags=re.IGNORECASE)  # Handle "eg,"

        # Step 2: Fix extra periods like `e.g..` or `e.g...,` and ensure proper punctuation
        corrected_text = re.sub(r'\.([.,])', r'\1', corrected_text)  # Removes an extra period before a comma or period
        corrected_text = re.sub(r'\.\.+', '.', corrected_text)  # Ensures only one period after e.g.

        # Step 3: Remove comma if e.g... is followed by it (e.g..., -> e.g.)
        corrected_text = re.sub(r'e\.g\.,', 'e.g.', corrected_text)

        # Step 4: Change e.g, to e.g.
        corrected_text = re.sub(r'e\.g,', 'e.g.', corrected_text)

        # Log changes if the run's text is updated
        if corrected_text != original_text:
            global_logs.append(
                f"[e.g. correction] Line {line_number}: {original_text.strip()} -> {corrected_text.strip()}"
            )

        # Update the run's text while preserving formatting
        run.text = corrected_text



# def enforce_ie_rule_with_logging(text):
#     lines = text.splitlines()
#     updated_lines = []
#     for line_number, line in enumerate(lines, start=1):
#         original_line = line

#         # Step 1: Match "ie" or "i.e." with optional surrounding spaces and punctuation
#         new_line = re.sub(r'\bie\b', 'i.e.', line, flags=re.IGNORECASE)  # Handle standalone "ie"
#         new_line = re.sub(r'\bie,\b', 'i.e.', new_line, flags=re.IGNORECASE)  # Handle "ie,"

#         # Step 2: Fix extra periods like `i.e..` or `i.e...,` and ensure proper punctuation
#         new_line = re.sub(r'\.([.,])', r'\1', new_line)  # Removes an extra period before a comma or period
#         new_line = re.sub(r'\.\.+', '.', new_line)  # Ensures only one period after i.e.

#         # Step 3: Remove comma if i.e... is followed by it (i.e..., -> i.e.)
#         new_line = re.sub(r'i\.e\.,', 'i.e.', new_line)
        
#         # Step 4: Change i.e, to i.e.
#         new_line = re.sub(r'i\.e,', 'i.e.', new_line)

#         # Log changes if the line is updated
#         if new_line != line:
#             global_logs.append(
#                 f"[i.e. correction] Line {line_number}: {line.strip()} -> {new_line.strip()}"
#             )
        
#         updated_lines.append(new_line)
#     return "\n".join(updated_lines)



def enforce_ie_rule_with_logging(paragraph, line_number):
    """
    Ensures consistent usage of 'i.e.' in the paragraph's runs and logs changes,
    preserving the formatting.
    
    :param paragraph: The paragraph object containing runs.
    :param line_number: The line number in the document for logging.
    """
    global global_logs  # Use a global log to record changes

    for run in paragraph.runs:
        original_text = run.text  # Store the original text of the run
        corrected_text = original_text

        # Step 1: Match "ie" or "i.e." with optional surrounding spaces and punctuation
        corrected_text = re.sub(r'\bie\b', 'i.e.', corrected_text, flags=re.IGNORECASE)  # Handle standalone "ie"
        corrected_text = re.sub(r'\bie,\b', 'i.e.', corrected_text, flags=re.IGNORECASE)  # Handle "ie,"

        # Step 2: Fix extra periods like `i.e..` or `i.e...,` and ensure proper punctuation
        corrected_text = re.sub(r'\.([.,])', r'\1', corrected_text)  # Removes an extra period before a comma or period
        corrected_text = re.sub(r'\.\.+', '.', corrected_text)  # Ensures only one period after i.e.

        # Step 3: Remove comma if i.e... is followed by it (i.e..., -> i.e.)
        corrected_text = re.sub(r'i\.e\.,', 'i.e.', corrected_text)
        
        # Step 4: Change i.e, to i.e.
        corrected_text = re.sub(r'i\.e,', 'i.e.', corrected_text)

        # Log changes if the run's text is updated
        if corrected_text != original_text:
            global_logs.append(
                f"[i.e. correction] Line {line_number}: {original_text.strip()} -> {corrected_text.strip()}"
            )

        # Update the run's text while preserving formatting
        run.text = corrected_text



# def standardize_etc(text):
#     lines = text.splitlines()
#     updated_lines = []
#     pattern = r'\b(e\.?tc|e\.t\.c|e\.t\.c\.|et\.?\s?c|et\s?c|etc\.?|etc|et cetera|etcetera|Etc\.?|Etc|‘and etc\.’|et\.?\s?cetera|etc\.?,?|etc\.?\.?|etc\,?\.?)\b'
    
#     for line_number, line in enumerate(lines, start=1):
#         original_line = line
        
#         # Replace all matches of "etc." variations with "etc."
#         new_line = re.sub(pattern, 'etc.', line, flags=re.IGNORECASE)
        
#         # Explicitly replace "etc.." with "etc."
#         new_line = re.sub(r'etc\.\.+', 'etc.', new_line)
        
#         # Explicitly replace "etc.." with "etc."
#         new_line = re.sub(r'etc\.\.+', 'etc.', new_line)
        
#         # Explicitly replace "etc.," with "etc."
#         new_line = re.sub(r'etc\.,', 'etc.', new_line)

#         # Log changes if the line is updated
#         if new_line != line:
#             global_logs.append(f"[etc. correction] Line {line_number}: {line.strip()} -> {new_line.strip()}")
        
#         updated_lines.append(new_line)
#     return "\n".join(updated_lines)


def standardize_etc(paragraph, line_number):
    """
    Standardizes all variations of "etc." in the paragraph's runs and logs changes, preserving formatting.
    
    :param paragraph: The paragraph object containing runs.
    :param line_number: The line number in the document for logging.
    """
    global global_logs  # Use a global log to record changes

    # Define the pattern for matching various "etc." variations
    pattern = r'\b(e\.?tc|e\.t\.c|e\.t\.c\.|et\.?\s?c|et\s?c|etc\.?|etc|et cetera|etcetera|Etc\.?|Etc|‘and etc\.’|et\.?\s?cetera|etc\.?,?|etc\.?\.?|etc\,?\.?)\b'

    for run in paragraph.runs:
        original_text = run.text  # Store the original text of the run
        corrected_text = original_text

        # Replace all matches of "etc." variations with "etc."
        corrected_text = re.sub(pattern, 'etc.', corrected_text, flags=re.IGNORECASE)
        
        # Explicitly replace "etc.." with "etc."
        corrected_text = re.sub(r'etc\.\.+', 'etc.', corrected_text)
        
        # Explicitly replace "etc.," with "etc."
        corrected_text = re.sub(r'etc\.,', 'etc.', corrected_text)

        # Log changes if the run's text is updated
        if corrected_text != original_text:
            global_logs.append(
                f"[etc. correction] Line {line_number}: {original_text.strip()} -> {corrected_text.strip()}"
            )

        # Update the run's text while preserving formatting
        run.text = corrected_text




# def adjust_ratios(text):
#     """
#     Ensures proper formatting of ratios with spaces around the colon (e.g., "1:2" -> "1 : 2").

#     Args:
#         text (str): Input text to process.

#     Returns:
#         str: Updated text.
#     """
#     global global_logs
#     def process_ratio(match):
#         original = match.group(0)
#         modified = f"{match.group(1)} : {match.group(2)}"
#         if original != modified:
#             line_number = text[:match.start()].count('\n') + 1
#             global_logs.append(
#                 f"[adjust_ratios] Line {line_number}: '{original}' -> '{modified}'"
#             )
#         return modified
#     return re.sub(r"(\d)\s*:\s*(\d)", process_ratio, text)



def adjust_ratios(paragraph, line_number):
    """
    Ensures proper formatting of ratios with spaces around the colon (e.g., "1:2" -> "1 : 2") in the paragraph's runs
    and logs changes, preserving formatting.

    Args:
        paragraph: The paragraph object containing runs.
        line_number: The line number in the document for logging.
    """
    global global_logs

    def process_ratio(match):
        original = match.group(0)
        modified = f"{match.group(1)} : {match.group(2)}"
        if original != modified:
            global_logs.append(
                f"[adjust_ratios] Line {line_number}: '{original}' -> '{modified}'"
            )
        return modified

    for run in paragraph.runs:
        original_text = run.text  # Store the original text of the run
        corrected_text = original_text
        corrected_text = re.sub(r"(\d)\s*:\s*(\d)", process_ratio, corrected_text)
        if corrected_text != original_text:
            run.text = corrected_text



def correct_chapter_numbering(text, chapter_counter):
    chapter_pattern = re.compile(r"(?i)\bchapter\s+((?:[IVXLCDM]+)|(?:[a-z]+)|(?:\d+))[:.]?\s")
    def replace_chapter_heading(match):
        chapter_content = match.group(1)
        if re.match(r"^[IVXLCDM]+$", chapter_content, re.IGNORECASE):
            chapter_number = roman.fromRoman(chapter_content.upper())
        elif re.match(r"^[a-z]+$", chapter_content, re.IGNORECASE):
            chapter_number = w2n.word_to_num(chapter_content.lower())
        else:
            chapter_number = int(chapter_content)
        return f"Chapter {chapter_number}: "
    return chapter_pattern.sub(replace_chapter_heading, text)



# def enforce_number_spelling_rule(text: str):
#     num_to_words = {
#         "1": "one", "2": "two", "3": "three", "4": "four", "5": "five",
#         "6": "six", "7": "seven", "8": "eight", "9": "nine"
#     }
#     units = r"(kg|g|mg|cm|mm|km|m|l|ml|%)"
#     sentences = re.split(r"(?<=[.!?])\s+", text)
#     updated_sentences = []
#     for sentence in sentences:
#         numbers = re.findall(r"\b\d+\b", sentence)
#         if any(int(num) >= 10 for num in numbers) and any(int(num) < 10 for num in numbers):
#             updated_sentences.append(sentence)
#             continue
#         def replace_number(match):
#             number = match.group()
#             if number in num_to_words:
#                 if re.search(rf"\b{number}\b\s+{units}", sentence):
#                     return number
#                 if re.search(rf"\b{number}-[a-zA-Z-]+", sentence):
#                     return num_to_words[number]
#                 return num_to_words[number]
#             return number
#         updated_sentence = re.sub(r"\b\d+\b", replace_number, sentence)
#         updated_sentences.append(updated_sentence)
#     return " ".join(updated_sentences)



def enforce_number_spelling_rule(paragraph, line_number):
    """
    Enforces a rule that numbers less than 10 are written as words, and logs the changes.
    This function processes each run of text in the paragraph to preserve formatting.

    :param paragraph: The paragraph object to process.
    :param line_number: The line number in the document for logging.
    """
    global global_logs

    num_to_words = {
        "1": "one", "2": "two", "3": "three", "4": "four", "5": "five",
        "6": "six", "7": "seven", "8": "eight", "9": "nine"
    }

    units = r"(kg|g|mg|cm|mm|km|m|l|ml|%)"

    def replace_number(match, sentence):
        number = match.group()
        if number in num_to_words:
            if re.search(rf"\b{number}\b\s+{units}", sentence):
                return number
            if re.search(rf"\b{number}-[a-zA-Z-]+", sentence):
                return num_to_words[number]
            return num_to_words[number]
        return number

    # Iterate through each run in the paragraph
    for run in paragraph.runs:
        original_text = run.text
        updated_text = original_text
        
        # Split the text of the run into sentences for processing
        sentences = re.split(r"(?<=[.!?])\s+", original_text)
        updated_sentences = []

        for sentence in sentences:
            numbers = re.findall(r"\b\d+\b", sentence)
            if any(int(num) >= 10 for num in numbers) and any(int(num) < 10 for num in numbers):
                updated_sentences.append(sentence)
                continue
            
            updated_sentence = re.sub(r"\b\d+\b", lambda match: replace_number(match, sentence), sentence)
            updated_sentences.append(updated_sentence)

        # Join the updated sentences back together
        updated_text = " ".join(updated_sentences)

        # If the text was modified, update the run and log the change
        if updated_text != original_text:
            global_logs.append(
                f"[number spelling] Line {line_number}: '{original_text.strip()}' -> '{updated_text.strip()}'"
            )
            run.text = updated_text




# Done
# [insert_thin_space_between_number_and_unit] Line 31: '5kg' -> '5 kg'
# def insert_thin_space_between_number_and_unit(text, line_number):
#     global global_logs
#     original_text = text
#     thin_space = '\u2009'
    
#     pattern = r"(\d+)(?=\s?[a-zA-Z]+)(?!\s?°)"

#     updated_text = text  # Initialize updated text to the original

#     matches = re.finditer(pattern, text)
#     for match in matches:
#         number = match.group(1)  # This is the number
#         unit_start = match.end()
#         unit = text[unit_start:].split()[0] 
        
#         original_word = number + unit
#         updated_word = number + thin_space + unit

#         updated_text = updated_text.replace(original_word, updated_word, 1)

#         global_logs.append(
#             f"[insert_thin_space_between_number_and_unit] Line {line_number}: '{original_word}' -> '{updated_word}'"
#         )
#     return updated_text


def insert_thin_space_between_number_and_unit(paragraph, line_number):
    """
    Inserts a thin space between numbers and units in a Word document paragraph,
    preserving the formatting (bold, italic, etc.) of the original text and logging the changes.

    :param paragraph: The paragraph object to process.
    :param line_number: The line number in the document for logging.
    """
    global global_logs
    thin_space = '\u2009'  # Unicode for thin space
    
    # Define the pattern to match numbers followed by units (e.g., "12kg", "5cm")
    pattern = r"(\d+)(?=\s?[a-zA-Z]+)(?!\s?°)"  # Matching number and unit but excluding degrees (°)

    # Iterate through each run in the paragraph
    for run in paragraph.runs:
        original_text = run.text
        updated_text = original_text
        
        matches = re.finditer(pattern, original_text)
        for match in matches:
            number = match.group(1)  # This is the number
            unit_start = match.end()
            unit = original_text[unit_start:].split()[0]  # Extract the unit that follows the number
            
            original_word = number + unit
            updated_word = number + thin_space + unit

            # Replace the original word with the updated word (insert thin space)
            updated_text = updated_text.replace(original_word, updated_word, 1)

            # Log the change if a modification occurs
            global_logs.append(
                f"[insert_thin_space_between_number_and_unit] Line {line_number}: '{original_word}' -> '{updated_word}'"
            )
        
        # If the text has been modified, update the run's text
        if updated_text != original_text:
            run.text = updated_text




# Done
# [format_dates] Line 5: '386 BCE' -> '386 bce'
# def format_dates(text, line_number):
#     global global_logs

#     def log_and_replace(pattern, replacement, text):
#         def replacer(match):
#             original = match.group(0)
#             updated = replacement(match)
#             if original != updated:
#                 global_logs.append(
#                     f"[format_dates] Line {line_number}: '{original}' -> '{updated}'"
#                 )
#             return updated
#         return re.sub(pattern, replacer, text)
#     text = log_and_replace(
#         r"\b(\d+)\s?(BCE|CE)\b",
#         lambda m: f"{m.group(1)} {m.group(2).lower()}",
#         text
#     )
#     text = log_and_replace(
#         r"\b(AD|BC)\.\b",
#         lambda m: f"{m.group(1)} ",
#         text
#     )
#     text = log_and_replace(
#         r"(\d+)\s?(BCE|CE|AD|BC)\b",
#         lambda m: f"{m.group(1)} {m.group(2)}",
#         text
#     )
#     return text



def format_dates(paragraph, line_number):
    """
    Formats dates (e.g., 'AD 2025', 'BCE 500') in a Word document paragraph, ensuring proper formatting
    and logs any changes made. This function modifies text while preserving the paragraph's formatting.
    
    :param paragraph: The paragraph object to process.
    :param line_number: The line number in the document for logging.
    """
    global global_logs

    # Define the helper function for logging and replacing text
    def log_and_replace(pattern, replacement, text):
        def replacer(match):
            original = match.group(0)
            updated = replacement(match)
            if original != updated:
                global_logs.append(
                    f"[format_dates] Line {line_number}: '{original}' -> '{updated}'"
                )
            return updated
        return re.sub(pattern, replacer, text)

    # Iterate through each run in the paragraph
    for run in paragraph.runs:
        original_text = run.text
        updated_text = original_text

        # Apply date formatting rules with logging
        updated_text = log_and_replace(
            r"\b(\d+)\s?(BCE|CE)\b",
            lambda m: f"{m.group(1)} {m.group(2).lower()}",
            updated_text
        )
        updated_text = log_and_replace(
            r"\b(AD|BC)\.\b",
            lambda m: f"{m.group(1)} ",
            updated_text
        )
        updated_text = log_and_replace(
            r"(\d+)\s?(BCE|CE|AD|BC)\b",
            lambda m: f"{m.group(1)} {m.group(2)}",
            updated_text
        )

        # If the text was modified, update the run's text
        if updated_text != original_text:
            run.text = updated_text




# Done
# [remove_space_between_degree_and_direction] Line 10: '52 °N' -> '52°N'
# def remove_space_between_degree_and_direction(text, line_number):
#     global global_logs
#     pattern = r"(\d+) \s*[º°]\s*(N|S|E|W)\b"
#     def log_replacement(match):
#         original_text = match.group(0)
#         updated_text = match.group(1) + "º" + match.group(2)
#         global_logs.append(
#             f"[remove_space_between_degree_and_direction] Line {line_number}: '{original_text}' -> '{updated_text}'"
#         )
#         return updated_text
#     updated_text = re.sub(pattern, log_replacement, text)
#     return updated_text



def remove_space_between_degree_and_direction(paragraph, line_number):
    """
    Removes the space between the degree symbol and direction (e.g., "30 ° N" -> "30ºN")
    while preserving the formatting in the paragraph (e.g., bold, italics) and logs changes.

    :param paragraph: The paragraph object to process.
    :param line_number: The line number in the document for logging.
    """
    global global_logs

    # Define the helper function for logging and replacing text
    def log_replacement(match):
        original_text = match.group(0)
        updated_text = match.group(1) + "º" + match.group(2)
        global_logs.append(
            f"[remove_space_between_degree_and_direction] Line {line_number}: '{original_text}' -> '{updated_text}'"
        )
        return updated_text

    # Iterate through each run in the paragraph
    for run in paragraph.runs:
        original_text = run.text
        updated_text = original_text

        # Apply the space removal rule with logging
        updated_text = re.sub(r"(\d+) \s*[º°]\s*(N|S|E|W)\b", log_replacement, updated_text)

        # If the text was modified, update the run's text
        if updated_text != original_text:
            run.text = updated_text



# Done
# km not Km; kg not Kg; l not L. (2.9)
# def enforce_lowercase_units(text, line_number):
#     global global_logs
#     unit_patterns = [
#         (r"(\d+)\s*(K)(m|g|l)", 'K', 'k'),
#         (r"(\d+)\s*(G)(m)", 'G', 'g'),
#         (r"(\d+)\s*(M)(g)", 'M', 'm'),
#         (r"(\d+)\s*(T)(g)", 'T', 't'),
#         (r"(\d+)\s*(L)\b", 'L', 'l'),
#         (r"(\d+)\s*(M)\b", 'M', 'm'),
#         (r"(\d+)\s*(kg|mg|g|cm|m|km|l|s|h|min)", r"\1 \2", None)
#     ]
#     updated_text = text
#     for pattern, original, updated in unit_patterns:
#         matches = re.finditer(pattern, updated_text)
#         for match in matches:
#             original_text = match.group(0)
#             if updated is not None:
#                 updated_text = updated_text.replace(original_text, original_text.replace(original, updated))
#                 global_logs.append(
#                     f"[enforce_lowercase_units] Line {line_number}: '{original_text}' -> '{original_text.replace(original, updated)}'"
#                 )
#             else:
#                 updated_text = updated_text.replace(original_text, f"{match.group(1)} {match.group(2)}")
#                 global_logs.append(
#                     f"[enforce_lowercase_units] Line {line_number}: '{original_text}' -> '{match.group(1)} {match.group(2)}'"
#                 )
#     return updated_text



def enforce_lowercase_units(paragraph, line_number):
    """
    Enforces lowercase formatting for units (e.g., "Kg" -> "kg") and ensures proper spacing between numbers and units.
    Also logs changes.

    :param paragraph: The paragraph object to process.
    :param line_number: The line number in the document for logging.
    """
    global global_logs

    # Define unit patterns and their conversions
    unit_patterns = [
        (r"(\d+)\s*(K)(m|g|l)", 'K', 'k'),
        (r"(\d+)\s*(G)(m)", 'G', 'g'),
        (r"(\d+)\s*(M)(g)", 'M', 'm'),
        (r"(\d+)\s*(T)(g)", 'T', 't'),
        (r"(\d+)\s*(L)\b", 'L', 'l'),
        (r"(\d+)\s*(M)\b", 'M', 'm'),
        (r"(\d+)\s*(kg|mg|g|cm|m|km|l|s|h|min)", r"\1 \2", None)
    ]
    
    # Iterate through each run in the paragraph
    for run in paragraph.runs:
        original_text = run.text
        updated_text = original_text

        # Apply the unit patterns and enforce lowercase units
        for pattern, original, updated in unit_patterns:
            updated_text = re.sub(pattern, lambda match: process_match(match, original, updated, line_number), updated_text)

        # If the text was modified, update the run's text
        if updated_text != original_text:
            run.text = updated_text

def process_match(match, original, updated, line_number):
    """
    Handles the replacement of units with lowercase versions, logging the changes.

    :param match: The match object found by the regular expression.
    :param original: The original unit to be replaced.
    :param updated: The updated lowercase unit.
    :param line_number: The line number in the document for logging.
    :return: The updated text after applying the change.
    """
    original_text = match.group(0)
    if updated is not None:
        updated_text = original_text.replace(original, updated)
        global_logs.append(
            f"[enforce_lowercase_units] Line {line_number}: '{original_text}' -> '{updated_text}'"
        )
    else:
        updated_text = f"{match.group(1)} {match.group(2)}"
        global_logs.append(
            f"[enforce_lowercase_units] Line {line_number}: '{original_text}' -> '{updated_text}'"
        )
    return updated_text




# Done
# [precede_decimal_with_zero] Line 22: '.76' -> '0.76'
# def precede_decimal_with_zero(text, line_number):
#     global global_logs
#     pattern = r"(?<!\d)(?<!\d\.)\.(\d+)"
#     def log_replacement(match):
#         original_text = match.group(0)
#         updated_text = "0." + match.group(1)
#         global_logs.append(
#             f"[precede_decimal_with_zero] Line {line_number}: '{original_text}' -> '{updated_text}'"
#         )
#         return updated_text
#     updated_text = re.sub(pattern, log_replacement, text)
#     return updated_text


def precede_decimal_with_zero(text, line_number, global_logs):
    """
    Ensures that decimals are properly formatted by preceding them with a zero if needed.
    Logs changes where applicable.

    :param text: The text to process.
    :param line_number: The line number for logging.
    :param global_logs: A list to store the logs of changes.
    :return: The updated text with proper decimal formatting.
    """
    # Define the pattern for detecting decimals without a leading zero
    pattern = r"(?<!\d)(?<!\d\.)\.(\d+)"

    # Function to handle replacement and logging
    def log_replacement(match):
        original_text = match.group(0)
        updated_text = "0." + match.group(1)

        # Log the replacement
        global_logs.append(
            f"[precede_decimal_with_zero] Line {line_number}: '{original_text}' -> '{updated_text}'"
        )
        return updated_text

    # Apply the pattern and replace decimals in the text
    updated_text = re.sub(pattern, log_replacement, text)

    return updated_text



# Done
# def adjust_terminal_punctuation_in_quotes(text):
#     text = re.sub(
#         r"([‘“])([^’”]*[?!])([’”])\.",
#         r"\1\2\3",
#         text
#     )
#     return text



def adjust_terminal_punctuation_in_quotes(text, line_number, global_logs):
    """
    Adjusts the punctuation inside quotes by removing the extra period after question marks or exclamation marks.
    Logs changes where applicable.

    :param text: The text to process.
    :param line_number: The line number for logging.
    :param global_logs: A list to store the logs of changes.
    :return: The updated text with correct punctuation in quotes.
    """
    # Define the pattern to match quotes with terminal punctuation
    pattern = r"([‘“])([^’”]*[?!])([’”])\."

    # Function to handle replacement and logging
    def log_replacement(match):
        original_text = match.group(0)
        updated_text = f"{match.group(1)}{match.group(2)}{match.group(3)}"

        # Log the replacement
        global_logs.append(
            f"[adjust_terminal_punctuation_in_quotes] Line {line_number}: '{original_text}' -> '{updated_text}'"
        )
        return updated_text

    # Apply the pattern and replace unnecessary periods in quotes
    updated_text = re.sub(pattern, log_replacement, text)

    return updated_text



# def enforce_serial_comma(text):
#     lines = text.splitlines()
#     updated_lines = []

#     for line_number, line in enumerate(lines, start=1):
#         original_line = line

#         # Add a comma before "and" or "or" in lists
#         new_line = re.sub(
#             r'([^,]+), ([^,]+) (or) ([^,]+)',
#             r'\1, \2, \3 \4',
#             line
#         )
#         # Explicitly handle cases where "or" does not get the serial comma
#         new_line = re.sub(
#             r'([^,]+), ([^,]+) (and) ([^,]+)',
#             r'\1, \2, \3 \4',
#             new_line
#         )
#         if new_line != line:
#             global_logs.append(f"[Serial comma correction] Line {line_number}: {line.strip()} -> {new_line.strip()}")
        
#         updated_lines.append(new_line)
#     return "\n".join(updated_lines)



def enforce_serial_comma(text, line_number, global_logs):
    """
    Adds a serial comma before "and" or "or" in lists, where necessary.
    Logs the changes where applicable.

    :param text: The text to process.
    :param line_number: The line number for logging.
    :param global_logs: A list to store the logs of changes.
    :return: The updated text with the serial comma applied.
    """
    # Define the pattern for adding serial commas before 'and' or 'or' in lists
    new_line = re.sub(
        r'([^,]+), ([^,]+) (or) ([^,]+)',
        r'\1, \2, \3 \4',
        text
    )

    # Explicitly handle cases for 'and'
    new_line = re.sub(
        r'([^,]+), ([^,]+) (and) ([^,]+)',
        r'\1, \2, \3 \4',
        new_line
    )

    # Log changes if there were any replacements
    if new_line != text:
        global_logs.append(f"[Serial comma correction] Line {line_number}: {text.strip()} -> {new_line.strip()}")
    
    return new_line



# def correct_possessive_names(text, line_number):
#     global global_logs
#     pattern_singular_possessive = r"\b([A-Za-z]+s)\b(?<!\bs')'"
#     matches_singular = re.finditer(pattern_singular_possessive, text)
#     updated_text = text
    
#     for match in matches_singular:
#         original_text = match.group(0)
#         updated_text_singular = match.group(1)[:-1] + "'s"
#         updated_text = updated_text.replace(original_text, updated_text_singular)
#         global_logs.append(
#             f"[correct_possessive_names] Line {line_number}: '{original_text}' -> '{updated_text_singular}'"
#         )
#     pattern_plural_possessive = r"\b([A-Za-z]+s)'\b"
#     matches_plural = re.finditer(pattern_plural_possessive, updated_text)
#     for match in matches_plural:
#         original_text = match.group(0)
#         updated_text_plural = match.group(1) + "'"
#         updated_text = updated_text.replace(original_text, updated_text_plural)
#         global_logs.append(
#             f"[correct_possessive_names] Line {line_number}: '{original_text}' -> '{updated_text_plural}'"
#         )
#     return updated_text



def correct_possessive_names(text, line_number, global_logs):
    """
    Corrects possessive forms of names. Ensures the correct format of possessive ('s) for singular and plural names.
    Logs changes where applicable.

    :param text: The text to process.
    :param line_number: The line number for logging.
    :param global_logs: A list to store the logs of changes.
    :return: The updated text with possessive names corrected.
    """
    # Handle singular possessive (e.g., "James'" to "James's")
    pattern_singular_possessive = r"\b([A-Za-z]+s)\b(?<!\bs')'"
    matches_singular = re.finditer(pattern_singular_possessive, text)
    updated_text = text
    
    for match in matches_singular:
        original_text = match.group(0)
        updated_text_singular = match.group(1)[:-1] + "'s"
        updated_text = updated_text.replace(original_text, updated_text_singular)
        global_logs.append(
            f"[correct_possessive_names] Line {line_number}: '{original_text}' -> '{updated_text_singular}'"
        )
    
    # Handle plural possessive (e.g., "James's" to "James'")
    pattern_plural_possessive = r"\b([A-Za-z]+s)'\b"
    matches_plural = re.finditer(pattern_plural_possessive, updated_text)
    for match in matches_plural:
        original_text = match.group(0)
        updated_text_plural = match.group(1) + "'"
        updated_text = updated_text.replace(original_text, updated_text_plural)
        global_logs.append(
            f"[correct_possessive_names] Line {line_number}: '{original_text}' -> '{updated_text_plural}'"
        )
    
    return updated_text





# Done
# http://www.PHi.com/authorguidelines not http://www.PHi.com/authorguidelines/
# def remove_concluding_slashes_from_urls(text, line_number):
#     global global_logs
#     pattern = r"(https?://[^\s/]+(?:/[^\s/]+)*)/"
#     matches = re.finditer(pattern, text)
#     updated_text = text
    
#     for match in matches:
#         original_text = match.group(0)
#         updated_text_url = match.group(1)  # URL without the concluding slash (e.g., "https://example.com")
#         updated_text = updated_text.replace(original_text, updated_text_url)
        
#         # Log the change
#         global_logs.append(
#             f"[remove_concluding_slashes_from_urls] Line {line_number}: '{original_text}' -> '{updated_text_url}'"
#         )    
#     return updated_text




def remove_concluding_slashes_from_urls(text, line_number, global_logs):
    """
    Removes concluding slashes from URLs. Ensures that URLs do not have a slash at the end.
    Logs changes where applicable.

    :param text: The text to process.
    :param line_number: The line number for logging.
    :param global_logs: A list to store the logs of changes.
    :return: The updated text with concluding slashes removed from URLs.
    """
    # Regex pattern to match URLs ending with a slash
    pattern = r"(https?://[^\s/]+(?:/[^\s/]+)*)/"
    
    # Find all matches of the pattern in the text
    matches = re.finditer(pattern, text)
    updated_text = text  # Initialize updated text as the original text
    
    for match in matches:
        original_text = match.group(0)  # Original URL with the concluding slash
        updated_text_url = match.group(1)  # URL without the concluding slash (e.g., "https://example.com")
        
        # Replace the original URL with the updated one
        updated_text = updated_text.replace(original_text, updated_text_url)
        
        # Log the change
        global_logs.append(
            f"[remove_concluding_slashes_from_urls] Line {line_number}: '{original_text}' -> '{updated_text_url}'"
        )
    
    return updated_text





# def clean_web_addresses(text):
#     """
#     Removes angle brackets around web addresses (e.g., "<http://example.com>" -> "http://example.com").
#     Args:
#         text (str): Input text to process.
#     Returns:
#         str: Updated text.
#     """
#     global global_logs
#     def process_web_address(match):
#         original = match.group(0)
#         modified = match.group(1)

#         if original != modified:
#             line_number = text[:match.start()].count('\n') + 1
#             global_logs.append(
#                 f"[clean_web_addresses] Line {line_number}: '{original}' -> '{modified}'"
#             )
#         return modified
#     return re.sub(r"<(https?://[^\s<>]+)>", process_web_address, text)



def clean_web_addresses(text, global_logs):
    """
    Removes angle brackets around web addresses (e.g., "<http://example.com>" -> "http://example.com").
    
    Args:
        text (str): Input text to process.
        global_logs (list): List to store the log of changes.

    Returns:
        str: Updated text with cleaned web addresses.
    """
    def process_web_address(match):
        original = match.group(0)
        modified = match.group(1)

        if original != modified:
            line_number = text[:match.start()].count('\n') + 1
            global_logs.append(
                f"[clean_web_addresses] Line {line_number}: '{original}' -> '{modified}'"
            )
        return modified

    return re.sub(r"<(https?://[^\s<>]+)>", process_web_address, text)




# def format_ellipses_in_series(text):
#     # Matches series like "x1, x2, ..., xn" and ensures the ellipsis has a comma and space after it.
#     text = re.sub(r"(\w+),\s*(\w+),\s*\.\.\.\s*(\w+)", r"\1, \2, …, \3", text)
#     return text



def format_ellipses_in_series(text, global_logs, line_number):
    """
    Matches series like "x1, x2, ..., xn" and ensures the ellipsis has a comma and space after it,
    and logs the changes made.

    Args:
        text (str): Input text to process.
        global_logs (list): List to store the log of changes.
        line_number (int): The line number of the current text.

    Returns:
        str: Updated text with correctly formatted ellipses in series.
    """
    def log_replacement(match):
        original_text = match.group(0)
        updated_text = f"{match.group(1)}, {match.group(2)}, …, {match.group(3)}"
        if original_text != updated_text:
            global_logs.append(
                f"[format_ellipses_in_series] Line {line_number}: '{original_text}' -> '{updated_text}'"
            )
        return updated_text

    return re.sub(r"(\w+),\s*(\w+),\s*\.\.\.\s*(\w+)", log_replacement, text)






# def format_chapter_title(text):
#     match = re.match(r"Chapter\s+([\dIVXLCDM]+)[\.:]\s*(.*)", text, re.IGNORECASE)
#     if match:
#         chapter_number = match.group(1)
#         chapter_title = match.group(2).rstrip('.')
#         words = chapter_title.split()
#         formatted_title = " ".join([
#             word.capitalize() if i == 0 or len(word) >= 4 else word.lower()
#             for i, word in enumerate(words)
#         ])
#         # print(formatted_title)
#         return f"{chapter_number}. {formatted_title}"
#     return text


def format_chapter_title(text, global_logs, line_number):
    """
    Formats chapter titles by ensuring the proper format for the chapter number and title,
    and logs the changes made.

    Args:
        text (str): Input text to process.
        global_logs (list): List to store the log of changes.
        line_number (int): The line number of the current text.

    Returns:
        str: Updated text with the correctly formatted chapter title.
    """
    match = re.match(r"Chapter\s+([\dIVXLCDM]+)[\.:]\s*(.*)", text, re.IGNORECASE)
    
    if match:
        chapter_number = match.group(1)
        chapter_title = match.group(2).rstrip('.')
        words = chapter_title.split()
        formatted_title = " ".join([
            word.capitalize() if i == 0 or len(word) >= 4 else word.lower()
            for i, word in enumerate(words)
        ])
        updated_text = f"{chapter_number}. {formatted_title}"

        # Log the change
        if updated_text != text:
            global_logs.append(
                f"[format_chapter_title] Line {line_number}: '{text.strip()}' -> '{updated_text.strip()}'"
            )
        
        return updated_text
    
    return text



# def format_titles_us_english_with_logging(text, doc_id):
#     global global_logs

#     titles = {
#         "doctor": "Dr.",
#         "mister": "Mr.",
#         "misses": "Mrs.",
#         "miss": "Miss.",
#         "ms": "Ms.",
#         "professor": "Professor",
#         "sir": "Sir",
#         "madam": "Madam",
#         "saint": "St",
#     }    
    
#     lines = text.splitlines()
#     updated_lines = []

#     for line_number, line in enumerate(lines, start=1):
#         original_line = line
#         for title, replacement in titles.items():
#             # Replace case-insensitive title with formatted title
#             new_line = re.sub(rf"\b{title}\b", replacement, line, flags=re.IGNORECASE)
#             if new_line != line:
#                 # Log the change to the global array
#                 global_logs.append(f"[shorten title] Line {line_number}: {title} -> {replacement}")
#                 line = new_line
#         updated_lines.append(line)

#     # Return the updated text
#     return "\n".join(updated_lines)


# def units_with_bracket(text, doc_id):
#     units = {
#         "s": "second",
#         "m": "meter",
#         "kg": "kilogram",
#         "A": "ampere",
#         "K": "kelvin",
#         "mol": "mole",
#         "cd": "candela"
#     }

#     used_units = set()
#     global global_logs

#     processed_lines = []
#     for line_num, line in enumerate(text.splitlines(), start=1):
#         def replace_unit(match):
#             number = match.group(1)
#             unit = match.group(2)
            
#             if unit in used_units:
#                 return match.group(0)
#             else:
#                 used_units.add(unit)
#                 full_form = units[unit]

#                 if unit != "mol" and not full_form.endswith("s"):
#                     full_form += "s"
#                 replacement = f"{number} {full_form} ({unit.lower()})"
#                 global_logs.append(
#                     f"Line {line_num}: {match.group(0)} -> {replacement}"
#                 )

#                 return replacement
#         pattern = r'\b(\d+)\s*(%s)\b' % '|'.join(re.escape(unit) for unit in units.keys())
#         processed_line = re.sub(pattern, replace_unit, line)
#         processed_lines.append(processed_line)
#     return "\n".join(processed_lines)


def format_titles_us_english_with_logging(text, doc_id):
    global global_logs

    titles = {
        "doctor": "Dr.",
        "mister": "Mr.",
        "misses": "Mrs.",
        "miss": "Miss.",
        "ms": "Ms.",
        "professor": "Professor",
        "sir": "Sir",
        "madam": "Madam",
        "saint": "St",
    }

    lines = text.splitlines()
    updated_lines = []

    for line_number, line in enumerate(lines, start=1):
        original_line = line
        for title, replacement in titles.items():
            # Replace case-insensitive title with formatted title
            new_line = re.sub(rf"\b{title}\b", replacement, line, flags=re.IGNORECASE)
            if new_line != line:
                # Log the change to the global array
                global_logs.append(f"[shorten title] Line {line_number}: '{title}' -> '{replacement}'")
                line = new_line
        updated_lines.append(line)

    # Return the updated text
    return "\n".join(updated_lines)


def units_with_bracket(text, doc_id):
    units = {
        "s": "second",
        "m": "meter",
        "kg": "kilogram",
        "A": "ampere",
        "K": "kelvin",
        "mol": "mole",
        "cd": "candela"
    }

    used_units = set()
    global global_logs

    processed_lines = []
    for line_num, line in enumerate(text.splitlines(), start=1):
        def replace_unit(match):
            number = match.group(1)
            unit = match.group(2)
            
            if unit in used_units:
                return match.group(0)
            else:
                used_units.add(unit)
                full_form = units[unit]

                if unit != "mol" and not full_form.endswith("s"):
                    full_form += "s"
                replacement = f"{number} {full_form} ({unit.lower()})"
                global_logs.append(
                    f"Line {line_num}: '{match.group(0)}' -> '{replacement}'"
                )

                return replacement
        pattern = r'\b(\d+)\s*(%s)\b' % '|'.join(re.escape(unit) for unit in units.keys())
        processed_line = re.sub(pattern, replace_unit, line)
        processed_lines.append(processed_line)
    return "\n".join(processed_lines)




# def correct_scientific_units_with_logging(text):
#     global global_logs
#     unit_symbols = ['kg', 'm', 's', 'A', 'K', 'mol', 'cd', 'Hz', 'N', 'Pa', 'J', 'W', 'C', 'V', 'F', 'Ω', 'ohm', 'S', 'T', 'H', 'lm', 'lx', 'Bq', 'Gy', 'Sv', 'kat']
#     pattern = rf"\b(\d+)\s*({'|'.join(re.escape(unit) for unit in unit_symbols)})\s*(s|'s|\.s)?\b"
#     lines = text.splitlines()
#     updated_lines = []
    
#     for line_number, line in enumerate(lines, start=1):
#         original_line = line
#         changes = []
#         new_line = re.sub(pattern, lambda m: f"{m.group(1)} {m.group(2)}", line)
                
#         if new_line != line:
#             for match in re.finditer(pattern, line):
#                 original = match.group(0)
#                 corrected = f"{match.group(1)} {match.group(2)}"
#                 if original != corrected:
#                     changes.append(f"'{original}' -> '{corrected}'")

#             if changes:
#                 global_logs.append(
#                     f"[unit correction] Line {line_number}: {', '.join(changes)}"
#                 )

#         updated_lines.append(new_line)
        
#     return "\n".join(updated_lines)




def correct_scientific_units_with_logging(text):
    global global_logs
    unit_symbols = ['kg', 'm', 's', 'A', 'K', 'mol', 'cd', 'Hz', 'N', 'Pa', 'J', 'W', 'C', 'V', 'F', 'Ω', 'ohm', 'S', 'T', 'H', 'lm', 'lx', 'Bq', 'Gy', 'Sv', 'kat']
    pattern = rf"\b(\d+)\s*({'|'.join(re.escape(unit) for unit in unit_symbols)})\s*(s|'s|\.s)?\b"
    
    lines = text.splitlines()
    updated_lines = []
    
    for line_number, line in enumerate(lines, start=1):
        original_line = line
        changes = []
        new_line = re.sub(pattern, lambda m: f"{m.group(1)} {m.group(2)}", line)
        
        if new_line != line:
            # Log all unit corrections in the line
            for match in re.finditer(pattern, line):
                original = match.group(0)
                corrected = f"{match.group(1)} {match.group(2)}"
                if original != corrected:
                    changes.append(f"'{original}' -> '{corrected}'")

            if changes:
                global_logs.append(
                    f"[unit correction] Line {line_number}: {', '.join(changes)}"
                )

        updated_lines.append(new_line)
        
    return "\n".join(updated_lines)



def write_to_log(doc_id):
    global global_logs

    output_dir = os.path.join('output', str(doc_id))
    os.makedirs(output_dir, exist_ok=True)
    log_file_path = os.path.join(output_dir, 'global_logs.txt')

    with open(log_file_path, 'w', encoding='utf-8') as log_file:
        log_file.write("\n".join(global_logs))

    global_logs = []






# twofold not two-fold hyphenate with numeral for numbers greater than nine, e.g. 10-fold. 
# def replace_fold_phrases(text):
#     """
#     Replaces phrases with '-fold' to ensure correct formatting based on the number preceding it.
#     Args:
#         text (str): Input text to process.
#     Returns:
#         str: Updated text.
#     """
#     global global_logs
#     def process_fold(match):
#         original = match.group(0)
#         num_str = match.group(1)
#         separator = match.group(2)
        
#         if separator != "-":
#             return original

#         try:
#             if num_str.isdigit():
#                 number = int(num_str)
#             else:
#                 number = w2n.word_to_num(num_str)

#             if number > 9:
#                 modified = f"{number}-fold"
#             else:
#                 modified = f"{num2words(number)}fold"

#             if original != modified:
#                 line_number = text[:match.start()].count('\n') + 1
#                 global_logs.append(
#                     f"[replace_fold_phrases] Line {line_number}: '{original}' -> '{modified}'"
#                 )
#             return modified
#         except ValueError:
#             return original

#     pattern = r"(\b\w+\b)(-?)fold"
#     updated_text = re.sub(pattern, process_fold, text)
#     return updated_text


def replace_fold_phrases(text):
    """
    Replaces phrases with '-fold' to ensure correct formatting based on the number preceding it.
    Args:
        text (str): Input text to process.
    Returns:
        str: Updated text.
    """
    global global_logs
    
    def process_fold(match):
        original = match.group(0)
        num_str = match.group(1)
        separator = match.group(2)
        
        if separator != "-":
            return original

        try:
            if num_str.isdigit():
                number = int(num_str)
            else:
                number = w2n.word_to_num(num_str)

            if number > 9:
                modified = f"{number}-fold"
            else:
                modified = f"{num2words(number)}fold"

            # Log the change if modified
            if original != modified:
                line_number = text[:match.start()].count('\n') + 1
                global_logs.append(
                    f"[replace_fold_phrases] Line {line_number}: '{original}' -> '{modified}'"
                )
            return modified
        except ValueError:
            return original

    pattern = r"(\b\w+\b)(-?)fold"
    updated_text = re.sub(pattern, process_fold, text)
    return updated_text




# def correct_preposition_usage(text):
#     """
#     Corrects preposition usage for date ranges (e.g., "from 2000-2010" -> "from 2000 to 2010").
#     Args:
#         text (str): Input text to process.
#     Returns:
#         str: Updated text.
#     """
#     global global_logs

#     def process_from_to(match):
#         original = match.group(0)
#         modified = f"from {match.group(1)} to {match.group(2)}"

#         if original != modified:
#             line_number = text[:match.start()].count('\n') + 1
#             global_logs.append(
#                 f"[correct_preposition_usage] Line {line_number}: '{original}' -> '{modified}'"
#             )

#         return modified

#     def process_between_and(match):
#         original = match.group(0)
#         modified = f"between {match.group(1)} and {match.group(2)}"

#         if original != modified:
#             line_number = text[:match.start()].count('\n') + 1
#             global_logs.append(
#                 f"[correct_preposition_usage] Line {line_number}: '{original}' -> '{modified}'"
#             )
#         return modified

#     text = re.sub(r"from (\d{4})[–-](\d{4})", process_from_to, text)
#     text = re.sub(r"between (\d{4})[–-](\d{4})", process_between_and, text)
#     return text



def correct_preposition_usage(text):
    """
    Corrects preposition usage for date ranges (e.g., "from 2000-2010" -> "from 2000 to 2010").
    Args:
        text (str): Input text to process.
    Returns:
        str: Updated text.
    """
    global global_logs

    def process_from_to(match):
        original = match.group(0)
        modified = f"from {match.group(1)} to {match.group(2)}"

        if original != modified:
            line_number = text[:match.start()].count('\n') + 1
            global_logs.append(
                f"[correct_preposition_usage] Line {line_number}: '{original}' -> '{modified}'"
            )

        return modified

    def process_between_and(match):
        original = match.group(0)
        modified = f"between {match.group(1)} and {match.group(2)}"

        if original != modified:
            line_number = text[:match.start()].count('\n') + 1
            global_logs.append(
                f"[correct_preposition_usage] Line {line_number}: '{original}' -> '{modified}'"
            )
        return modified

    # Correct 'from-to' date ranges
    text = re.sub(r"from (\d{4})[–-](\d{4})", process_from_to, text)

    # Correct 'between-and' date ranges
    text = re.sub(r"between (\d{4})[–-](\d{4})", process_between_and, text)

    return text






# def correct_scientific_unit_symbols(text):
#     """
#     Ensures proper capitalization of units derived from proper names (e.g., J, Hz, W, N) only when preceded by a number.

#     Args:
#         text (str): Input text to process.

#     Returns:
#         str: Updated text.
#     """
#     global global_logs

#     units = {
#         "j": "J",
#         "hz": "Hz",
#         "w": "W",
#         "n": "N",
#         "pa": "Pa",
#         "v": "V",
#         "a": "A",
#         "c": "C",
#         "lm": "lm",
#         "lx": "lx",
#         "t": "T",
#         "ohm": "Ω",
#         "s": "S",
#         "k": "K",
#         "cd": "cd",
#         "mol": "mol",
#         "rad": "rad",
#         "sr": "sr"
#     }

#     def process_unit(match):
#         original = match.group(0)
#         unit = match.group(2).lower()  # Capture the unit (second group)
#         modified = f"{match.group(1)}{units.get(unit, match.group(2))}"  # Replace with capitalized unit if in dictionary

#         if original != modified:
#             line_number = text[:match.start()].count('\n') + 1  # Calculate the line number
#             global_logs.append(
#                 f"[correct_scientific_unit_symbols] Line {line_number}: '{original}' -> '{modified}'"
#             )

#         return modified

#     # Regex to match a number followed by optional space and a unit
#     pattern = r"\b(\d+\s*)(%s)\b" % "|".join(re.escape(unit) for unit in units.keys())
#     updated_text = re.sub(pattern, process_unit, text, flags=re.IGNORECASE)
#     return updated_text



def correct_scientific_unit_symbols(text):
    """
    Ensures proper capitalization of units derived from proper names (e.g., J, Hz, W, N) only when preceded by a number.

    Args:
        text (str): Input text to process.

    Returns:
        str: Updated text.
    """
    global global_logs

    # Dictionary of units with proper capitalization
    units = {
        "j": "J",
        "hz": "Hz",
        "w": "W",
        "n": "N",
        "pa": "Pa",
        "v": "V",
        "a": "A",
        "c": "C",
        "lm": "lm",
        "lx": "lx",
        "t": "T",
        "ohm": "Ω",
        "s": "S",
        "k": "K",
        "cd": "cd",
        "mol": "mol",
        "rad": "rad",
        "sr": "sr"
    }

    def process_unit(match):
        original = match.group(0)
        unit = match.group(2).lower()  # Capture the unit (second group)
        modified = f"{match.group(1)}{units.get(unit, match.group(2))}"  # Replace with capitalized unit if in dictionary

        if original != modified:
            line_number = text[:match.start()].count('\n') + 1  # Calculate the line number
            global_logs.append(
                f"[correct_scientific_unit_symbols] Line {line_number}: '{original}' -> '{modified}'"
            )

        return modified

    # Regex to match a number followed by optional space and a unit
    pattern = r"\b(\d+\s*)(%s)\b" % "|".join(re.escape(unit) for unit in units.keys())
    updated_text = re.sub(pattern, process_unit, text, flags=re.IGNORECASE)
    return updated_text




# def remove_quotation(text: str):
#     """
#     Removes single quotation marks (') following capitalized words.
#     Args:
#         text (str): Input text to process.
#     Returns:
#         str: Updated text.
#     """
#     global global_logs

#     pattern = r"([A-Z]+)'"

#     def process_quotation_removal(match):
#         original = match.group(0)
#         modified = f"{match.group(1)}"

#         if original != modified:
#             line_number = text[:match.start()].count('\n') + 1
#             global_logs.append(
#                 f"[remove_quotation] Line {line_number}: '{original}' -> '{modified}'"
#             )
#         return modified
#     para_text = re.sub(pattern, process_quotation_removal, text)
#     return para_text



def remove_quotation(text: str):
    """
    Removes single quotation marks (') following capitalized words.
    Args:
        text (str): Input text to process.
    Returns:
        str: Updated text.
    """
    global global_logs

    pattern = r"([A-Z]+)'"

    def process_quotation_removal(match):
        original = match.group(0)
        modified = f"{match.group(1)}"

        if original != modified:
            line_number = text[:match.start()].count('\n') + 1
            global_logs.append(
                f"[remove_quotation] Line {line_number}: '{original}' -> '{modified}'"
            )
        return modified

    para_text = re.sub(pattern, process_quotation_removal, text)
    return para_text



# def remove_and(text: str):
#     """
#     Replaces 'and' between two capitalized words with an ampersand (&).
#     Args:
#         text (str): Input text to process.
#     Returns:
#         str: Updated text.
#     """
#     global global_logs

#     # Regex pattern to match "and" between two capitalized words
#     pattern = r'([A-Z][a-z]+)\s+and\s+([A-Z][a-z]+)'

#     def process_and_replacement(match):
#         original = match.group(0)
#         modified = f"{match.group(1)} & {match.group(2)}"

#         if original != modified:
#             line_number = text[:match.start()].count('\n') + 1
#             global_logs.append(
#                 f"[remove_and] Line {line_number}: '{original}' -> '{modified}'"
#             )
#         return modified
#     text = re.sub(pattern, process_and_replacement, text)
#     return text


def remove_and(text: str):
    """
    Replaces 'and' between two capitalized words with an ampersand (&).
    Args:
        text (str): Input text to process.
    Returns:
        str: Updated text.
    """
    global global_logs

    # Regex pattern to match "and" between two capitalized words
    pattern = r'([A-Z][a-z]+)\s+and\s+([A-Z][a-z]+)'

    def process_and_replacement(match):
        original = match.group(0)
        modified = f"{match.group(1)} & {match.group(2)}"

        if original != modified:
            line_number = text[:match.start()].count('\n') + 1
            global_logs.append(
                f"[remove_and] Line {line_number}: '{original}' -> '{modified}'"
            )
        return modified

    text = re.sub(pattern, process_and_replacement, text)
    return text



# def correct_units_in_ranges_with_logging(text):
#     global global_logs

#     # List of valid unit symbols
#     unit_symbols = ['cm', 'm', 'kg', 's', 'A', 'K', 'mol', 'cd', '%']

#     # Regex patterns
#     range_pattern = rf"\b(\d+)\s*({'|'.join(re.escape(unit) for unit in unit_symbols)})\s*(to|-|–|—)\s*(\d+)\s*\2\b"
#     thin_space_pattern = rf"\b(\d+)\s+({'|'.join(re.escape(unit) for unit in unit_symbols)})\b"

#     lines = text.splitlines()
#     updated_lines = []

#     for line_number, line in enumerate(lines, start=1):
#         original_line = line

#         # Correct repeated units in ranges
#         new_line = re.sub(
#             range_pattern,
#             lambda m: f"{m.group(1)} {m.group(3)} {m.group(4)} {m.group(2)}",
#             line
#         )

#         # Add thin space between value and unit (except %)
#         new_line = re.sub(
#             thin_space_pattern,
#             lambda m: f"{m.group(1)} {m.group(2)}" if m.group(2) != "%" else f"{m.group(1)}{m.group(2)}",
#             new_line
#         )

#         # Log changes if any
#         if new_line != line:
#             change_details = f"{line.strip()} -> {new_line.strip()}"
#             global_logs.append(f"Line {line_number}: {change_details}")
#             line = new_line

#         updated_lines.append(line)

#     # Return the updated text
#     return "\n".join(updated_lines)



def correct_units_in_ranges_with_logging(text):
    global global_logs

    # List of valid unit symbols
    unit_symbols = ['cm', 'm', 'kg', 's', 'A', 'K', 'mol', 'cd', '%']

    # Regex patterns
    range_pattern = rf"\b(\d+)\s*({'|'.join(re.escape(unit) for unit in unit_symbols)})\s*(to|-|–|—)\s*(\d+)\s*\2\b"
    thin_space_pattern = rf"\b(\d+)\s+({'|'.join(re.escape(unit) for unit in unit_symbols)})\b"

    lines = text.splitlines()
    updated_lines = []

    for line_number, line in enumerate(lines, start=1):
        original_line = line

        # Correct repeated units in ranges
        new_line = re.sub(
            range_pattern,
            lambda m: f"{m.group(1)} {m.group(3)} {m.group(4)} {m.group(2)}",
            line
        )

        # Add thin space between value and unit (except %)
        new_line = re.sub(
            thin_space_pattern,
            lambda m: f"{m.group(1)} {m.group(2)}" if m.group(2) != "%" else f"{m.group(1)}{m.group(2)}",
            new_line
        )

        # Log changes if any
        if new_line != line:
            change_details = f"'{line.strip()}' -> '{new_line.strip()}'"
            global_logs.append(f"[correct_units_in_ranges_with_logging] Line {line_number}: {change_details}")
            line = new_line

        updated_lines.append(line)

    # Return the updated text
    return "\n".join(updated_lines)




# def correct_unit_spacing(text):
#     """
#     Corrects spacing between numbers and units (e.g., "100 MB" -> "100MB").
#     Args:
#         text (str): Input text to process.
#     Returns:
#         str: Updated text.
#     """
#     global global_logs

#     units = ["Hz", "KHz", "MHz", "GHz", "kB", "MB", "GB", "TB"]
#     pattern = r"(\d+)\s+(" + "|".join(units) + r")"

#     def process_spacing(match):
#         original = match.group(0)
#         modified = f"{match.group(1)}{match.group(2)}"

#         if original != modified:
#             line_number = text[:match.start()].count('\n') + 1
#             global_logs.append(
#                 f"[correct_unit_spacing] Line {line_number}: '{original}' -> '{modified}'"
#             )
#         return modified
#     corrected_text = re.sub(pattern, process_spacing, text)
#     return corrected_text


def correct_unit_spacing(text):
    """
    Corrects spacing between numbers and units (e.g., "100 MB" -> "100MB").
    Args:
        text (str): Input text to process.
    Returns:
        str: Updated text.
    """
    global global_logs

    units = ["Hz", "KHz", "MHz", "GHz", "kB", "MB", "GB", "TB"]
    pattern = r"(\d+)\s+(" + "|".join(units) + r")"

    def process_spacing(match):
        original = match.group(0)
        modified = f"{match.group(1)}{match.group(2)}"

        if original != modified:
            line_number = text[:match.start()].count('\n') + 1
            global_logs.append(
                f"[correct_unit_spacing] Line {line_number}: '{original}' -> '{modified}'"
            )
        return modified

    corrected_text = re.sub(pattern, process_spacing, text)
    return corrected_text



# def apply_quotation_punctuation_rule(text: str):
#     """
#     Adjusts the placement of punctuation marks within single quotes.
#     Args:
#         text (str): Input text to process.
#     Returns:
#         str: Updated text.
#     """
#     global global_logs

#     pattern = r"‘(.*?)’([!?])"

#     def process_quotation_punctuation(match):
#         original = match.group(0)
#         modified = f"‘{match.group(1)}{match.group(2)}’"

#         if original != modified:
#             line_number = text[:match.start()].count('\n') + 1
#             global_logs.append(
#                 f"[apply_quotation_punctuation_rule] Line {line_number}: '{original}' -> '{modified}'"
#             )
#         return modified
#     updated_text = re.sub(pattern, process_quotation_punctuation, text)
#     return updated_text



def apply_quotation_punctuation_rule(text: str):
    """
    Adjusts the placement of punctuation marks within single quotes.
    Args:
        text (str): Input text to process.
    Returns:
        str: Updated text.
    """
    global global_logs

    pattern = r"‘(.*?)’([!?])"

    def process_quotation_punctuation(match):
        original = match.group(0)
        modified = f"‘{match.group(1)}{match.group(2)}’"

        if original != modified:
            line_number = text[:match.start()].count('\n') + 1
            global_logs.append(
                f"[apply_quotation_punctuation_rule] Line {line_number}: '{original}' -> '{modified}'"
            )
        return modified

    updated_text = re.sub(pattern, process_quotation_punctuation, text)
    return updated_text



# def enforce_dnase_rule(text: str):
#     """
#     Enforces the correct capitalization for 'DNase'.
#     Args:
#         text (str): Input text to process.
#     Returns:
#         str: Updated text.
#     """
#     global global_logs

#     pattern = r"\bDNAse\b"

#     def process_dnase_replacement(match):
#         original = match.group(0)
#         modified = "DNase"

#         if original != modified:
#             line_number = text[:match.start()].count('\n') + 1
#             global_logs.append(
#                 f"[enforce_dnase_rule] Line {line_number}: '{original}' -> '{modified}'"
#             )
#         return modified

#     updated_text = re.sub(pattern, process_dnase_replacement, text)
#     return updated_text



def enforce_dnase_rule(text: str):
    """
    Enforces the correct capitalization for 'DNase'.
    Args:
        text (str): Input text to process.
    Returns:
        str: Updated text.
    """
    global global_logs

    pattern = r"\bDNAse\b"

    def process_dnase_replacement(match):
        original = match.group(0)
        modified = "DNase"

        if original != modified:
            line_number = text[:match.start()].count('\n') + 1
            global_logs.append(
                f"[enforce_dnase_rule] Line {line_number}: '{original}' -> '{modified}'"
            )
        return modified

    updated_text = re.sub(pattern, process_dnase_replacement, text)
    return updated_text




# def apply_remove_italics_see_rule(text):
#     return text.replace('*see*', 'see')



def apply_remove_italics_see_rule(text: str):
    """
    Replaces the italics formatting for the word 'see' with plain text.
    Args:
        text (str): Input text to process.
    Returns:
        str: Updated text.
    """
    global global_logs

    pattern = r"\*see\*"

    def process_see_replacement(match):
        original = match.group(0)
        modified = "see"

        if original != modified:
            line_number = text[:match.start()].count('\n') + 1
            global_logs.append(
                f"[apply_remove_italics_see_rule] Line {line_number}: '{original}' -> '{modified}'"
            )
        return modified

    updated_text = re.sub(pattern, process_see_replacement, text)
    return updated_text



# def replace_ampersand(text):
#     global global_logs
#     def replacement(match):
#         left, right = match.group(1), match.group(2)
#         original = match.group(0)
#         line_number = text[:match.start()].count('\n') + 1
#         if left[0].isupper() and right[0].isupper():
#             return original

#         modified = left + ' and ' + right
#         global_logs.append(
#             f"[replace_ampersand] Line {line_number}: '{original}' -> '{modified}'"
#         )
#         return modified
#     return re.sub(r'(?m)(\w+)\s*&\s*(\w+)', replacement, text)



def replace_ampersand(text):
    """
    Replaces ampersands (&) with 'and' between two words, unless both words start with capital letters.
    Args:
        text (str): Input text to process.
    Returns:
        str: Updated text.
    """
    global global_logs

    def replacement(match):
        left, right = match.group(1), match.group(2)
        original = match.group(0)
        line_number = text[:match.start()].count('\n') + 1
        
        # Check if both left and right words start with capital letters
        if left[0].isupper() and right[0].isupper():
            return original

        # Replace & with 'and'
        modified = left + ' and ' + right
        
        # Log the change
        global_logs.append(
            f"[replace_ampersand] Line {line_number}: '{original}' -> '{modified}'"
        )
        return modified

    # Regex to match and replace ampersands
    return re.sub(r'(?m)(\w+)\s*&\s*(\w+)', replacement, text)



# def rename_section(text):
#     # Replace all occurrences of the § symbol with 'Section'
#     return re.sub(r'§', 'Section', text)



def rename_section(text):
    """
    Replaces all occurrences of the § symbol with 'Section' and logs the changes.
    Args:
        text (str): Input text to process.
    Returns:
        str: Updated text.
    """
    global global_logs
    pattern = r'§'
    def replacement(match):
        original = match.group(0)
        modified = 'Section'
        line_number = text[:match.start()].count('\n') + 1

        # Log the change
        global_logs.append(
            f"[rename_section] Line {line_number}: '{original}' -> '{modified}'"
        )
        return modified
    return re.sub(pattern, replacement, text)




# def process_url_add_http(text):
#     """
#     Adjusts URLs in the input text based on the given rules:
#     1. If a URL starts with 'www.' but doesn't have 'http://', prepend 'http://'.
#     2. If a URL already starts with 'http://', remove 'http://'.

#     Args:
#         text (str): The input text containing URLs.

#     Returns:
#         str: The modified text with URLs adjusted.
#     """
#     global global_logs

#     def add_http_prefix(match):
#         original = match.group(0)
#         modified = f"http://{match.group(1)}"

#         if original != modified:
#             line_number = text[:match.start()].count('\n') + 1
#             global_logs.append(
#                 f"[process_url_add_http] Line {line_number}: '{original}' -> '{modified}'"
#             )

#         return modified

#     def remove_http_prefix(match):
#         original = match.group(0)
#         modified = match.group(1)

#         if original != modified:
#             line_number = text[:match.start()].count('\n') + 1  # Calculate the line number
#             global_logs.append(
#                 f"[process_url_add_http] Line {line_number}: '{original}' -> '{modified}'"
#             )

#         return modified

#     text = re.sub(r"\bhttp://(www\.\S+)", remove_http_prefix, text)
#     text = re.sub(r"\b(www\.\S+)", add_http_prefix, text)
#     return text



def process_url_add_http(text):
    """
    Adjusts URLs in the input text based on the given rules:
    1. If a URL starts with 'www.' but doesn't have 'http://', prepend 'http://'.
    2. If a URL already starts with 'http://', remove 'http://'.

    Args:
        text (str): The input text containing URLs.

    Returns:
        str: The modified text with URLs adjusted.
    """
    global global_logs

    def add_http_prefix(match):
        original = match.group(0)
        modified = f"http://{match.group(1)}"

        if original != modified:
            line_number = text[:match.start()].count('\n') + 1
            global_logs.append(
                f"[process_url_add_http] Line {line_number}: '{original}' -> '{modified}'"
            )

        return modified

    def remove_http_prefix(match):
        original = match.group(0)
        modified = match.group(1)

        if original != modified:
            line_number = text[:match.start()].count('\n') + 1  # Calculate the line number
            global_logs.append(
                f"[process_url_add_http] Line {line_number}: '{original}' -> '{modified}'"
            )

        return modified

    # Replace URLs by removing 'http://' or adding 'http://' as necessary
    text = re.sub(r"\bhttp://(www\.\S+)", remove_http_prefix, text)
    text = re.sub(r"\b(www\.\S+)", add_http_prefix, text)
    
    return text




# def process_url_remove_http(url):
#     """
#     Removes 'http://' from a URL if there is no path, parameters, query, or fragment.
#     Args:
#         url (str): The input URL to process.
#     Returns:
#         str: The modified URL with 'http://' removed if applicable.
#     """
#     global global_logs

#     parsed = urlparse(url)
#     original = url

#     if parsed.scheme == "http" and not (parsed.path or parsed.params or parsed.query or parsed.fragment):
#         modified = parsed.netloc

#         if original != modified:
#             line_number = 1
#             global_logs.append(
#                 f"[process_url_remove_http] Line {line_number}: '{original}' -> '{modified}'"
#             )
#         return modified
#     return url




def process_url_remove_http(text):
    """
    Removes 'http://' from a URL if , or fragment.
    Args:
        text (str): The input text containing URLs.
    Returns:
        str: The modified text with 'http://' removed from URLs if applicable.
    """
    global global_logs

    def remove_http_prefix(match):
        original = match.group(0)
        url = match.group(1)
        
        parsed = urlparse(url)

        if parsed.scheme == "http" and not (parsed.path or parsed.params or parsed.query or parsed.fragment):
            modified = parsed.netloc

            if original != modified:
                line_number = text[:match.start()].count('\n') + 1
                global_logs.append(
                    f"[process_url_remove_http] Line {line_number}: '{original}' -> '{modified}'"
                )
            return modified
        return original
    pattern = r"\bhttp://([^\s]+)\b"
    updated_text = re.sub(pattern, remove_http_prefix, text)    
    return updated_text



# def process_symbols_mark(text, line_number, symbols=["®", "™", "©", "℗", "℠"]):
#     """
#     Ensures symbols like ®, ™, etc., appear only the first time in the text.
#     Updates the global_log with changes, including line number, original text, and updated text.
#     """
#     original_text = text
#     symbol_set = set()
#     global global_logs
    
#     for symbol in symbols:
#         occurrences = list(re.finditer(re.escape(symbol), text))
#         if occurrences:
#             first_occurrence = occurrences[0].start()
#             # Replace all occurrences after the first one
#             text = (
#                 text[:first_occurrence + 1]
#                 + re.sub(re.escape(symbol), "", text[first_occurrence + 1:])
#             )
#             symbol_set.add(symbol)

#     # Log changes if the text was modified
#     if original_text != text:
#         global_logs.append(
#             f"[process_symbols_in_doc] Line {line_number}: '{original_text}' -> '{text}'"
#         )
#     return text



def process_symbols_mark(text, line_number, symbols=["®", "™", "©", "℗", "℠"]):
    """
    Ensures symbols like ®, ™, etc., appear only the first time in the text.
    Updates the global_log with changes, including line number, original text, and updated text.
    
    Args:
        text (str): Input text to process.
        line_number (int): The line number where the change occurred.
        symbols (list): List of symbols to be processed (default includes common trademarks and copyright symbols).
    
    Returns:
        str: Updated text with duplicate symbols removed after the first occurrence.
    """
    global global_logs
    
    original_text = text
    symbol_set = set()

    for symbol in symbols:
        occurrences = list(re.finditer(re.escape(symbol), text))
        
        if occurrences:
            first_occurrence = occurrences[0].start()
            # Replace all occurrences after the first one
            text = (
                text[:first_occurrence + 1]
                + re.sub(re.escape(symbol), "", text[first_occurrence + 1:])
            )
            symbol_set.add(symbol)

    # Log changes if the text was modified
    if original_text != text:
        global_logs.append(
            f"[process_symbols_mark] Line {line_number}: '{original_text.strip()}' -> '{text.strip()}'"
        )    
    return text


# def remove_commas_from_numbers(text, line_number):
#     """
#     Removes commas from numerical values in the text.
#     Updates the global_log with the specific changes, including line number and changes made.
#     """
#     original_text = text
#     changes = []
#     global global_logs

#     pattern = r'\b\d{1,3}(,\d{3})+\b'

#     def replacer(match):
#         original_number = match.group(0)
#         updated_number = original_number.replace(",", " ")
#         changes.append((original_number, updated_number))
#         return updated_number

#     # Replace numbers with commas in the text
#     text = re.sub(pattern, replacer, text)

#     # Log individual changes
#     for original, updated in changes:
#         global_logs.append(
#             f"[process_symbols_in_doc] Line {line_number}: '{original}' -> '{updated}'"
#         )
#     return text

def remove_commas_from_numbers(runs, line_number):
    """
    Removes commas from numerical values in the runs of a paragraph.
    Updates the global_log with the specific changes, including line number and changes made.
    """
    import re

    global global_logs
    changes = []

    # Regex to match numbers with commas (e.g., 1,000 or 20,000)
    pattern = r'\b\d{1,3}(,\d{3})+\b'

    # Iterate through runs to handle text while preserving formatting
    for run in runs:
        original_text = run.text

        # Function to replace matched patterns while logging changes
        def replacer(match):
            original_number = match.group(0)  # Match the original number
            updated_number = original_number.replace(",", "")  # Remove commas
            changes.append((original_number, updated_number))  # Log the change
            return updated_number

        # Replace numbers with commas in the current run
        updated_text = re.sub(pattern, replacer, original_text)

        # If the text was modified, update the run text
        if original_text != updated_text:
            run.text = updated_text

    # Log individual changes
    for original, updated in changes:
        global_logs.append(
            f"[remove_commas_from_numbers] Line {line_number}: '{original}' -> '{updated}'"
        )




# def remove_spaces_from_four_digit_numbers(text, line_number):
#     """
#     Removes spaces from four-digit numerals in the text.
#     Updates the global_log with specific changes, including line number and changes made.
#     """
#     original_text = text
#     changes = []
#     global global_logs

#     pattern = r'\b\d\s\d{3}\b'

#     def replacer(match):
#         original_number = match.group(0)  # Match the original number
#         updated_number = original_number.replace(" ", "")  # Remove spaces
#         changes.append((original_number, updated_number))  # Log the change
#         return updated_number

#     text = re.sub(pattern, replacer, text)

#     for original, updated in changes:
#         global_logs.append(
#             f"[process_symbols_in_doc] Line {line_number}: '{original}' -> '{updated}'"
#         )
#     return text

def remove_spaces_from_four_digit_numbers(runs, line_number):
    """
    Removes spaces from four-digit numerals in the text of runs.
    Updates the global_log with specific changes, including line number and changes made.
    """
    import re

    global global_logs
    changes = []

    pattern = r'\b\d\s\d{3}\b'

    for run in runs:
        original_text = run.text
        
        def replacer(match):
            original_number = match.group(0)  # Match the original number
            updated_number = original_number.replace(" ", "")  # Remove spaces
            changes.append((original_number, updated_number))  # Log the change
            return updated_number

        # Replace the text in the current run while keeping its formatting
        updated_text = re.sub(pattern, replacer, original_text)
        run.text = updated_text

    for original, updated in changes:
        global_logs.append(
            f"[process_symbols_in_doc] Line {line_number}: '{original}' -> '{updated}'"
        )

    return runs



# def set_latinisms_to_roman_in_runs(paragraph_text, line_number, latinisms=None):
#     """
#     Converts specific Latinisms from italic to roman text in a string of text.
#     Logs changes to the global_log, including line number and original italicized Latinism.
#     """
#     if latinisms is None:
#         latinisms = [
#             "i.e.", "e.g.", "via", "vice versa", "etc.", "a posteriori", 
#             "a priori", "et al.", "cf.", "c."
#         ]
    
#     changes = []
#     global global_logs

#     # Process the text, and for each Latinism, replace its italics if needed
#     for lat in latinisms:
#         if lat in paragraph_text:
#             changes.append(lat)  # Log the Latinism that was changed

#     # for changed in changes:
#     #     global_logs.append(
#     #         f"[process_symbols_in_doc] Line {line_number}: '{changed}' -> '{changed}'"
#     #     )

#     return paragraph_text


def set_latinisms_to_roman_in_runs(runs, line_number, latinisms=None):
    """
    Converts specific Latinisms from italic to roman text in runs.
    Logs changes to the global_log, including line number and original italicized Latinism.
    """
    if latinisms is None:
        latinisms = [
            "i.e.", "e.g.", "via", "vice versa", "etc.", "a posteriori", 
            "a priori", "et al.", "cf.", "c."
        ]
    
    changes = []
    global global_logs

    # Process each run to replace italics for specified Latinisms
    for run in runs:
        for lat in latinisms:
            if lat in run.text:
                if run.italic:  # Check if the text is italic
                    run.italic = False  # Set to roman (non-italic)
                    changes.append(lat)  # Log the Latinism that was changed

    for changed in changes:
        global_logs.append(
            f"[process_symbols_in_doc] Line {line_number}: '{changed}' changed from italic to roman"
        )

    return runs




# def convert_decimal_to_baseline(paragraph_text, line_number):
#     """
#     Converts any non-standard decimal separator (•) to a standard decimal point (.)
#     only when both sides are numeric.
#     Logs the changes to the global_log, including line number and the change.
#     """
#     changes = []
#     global global_logs
#     # Regular expression to find '•' between numbers
#     pattern = r'(?<=\d)\xB7(?=\d)'

#     # Find all occurrences of '•' that are between digits and replace with '.'
#     matches = re.findall(pattern, paragraph_text)
#     if matches:
#         original_text = paragraph_text
#         updated_text = re.sub(pattern, '.', paragraph_text)  # Replace '•' with '.'
#         changes.append((original_text, updated_text))

#     for original, updated in changes:
#         global_logs.append(
#             f"[convert_decimal_to_baseline] Line {line_number}: '{original}' -> '{updated}'"
#         )
#     return updated_text if changes else paragraph_text


def convert_decimal_to_baseline(runs, line_number):
    """
    Converts any non-standard decimal separator (•) to a standard decimal point (.)
    only when both sides are numeric.
    Logs the changes to the global_log, including line number and the change.
    """
    import re

    changes = []
    global global_logs
    # Regular expression to find '•' between numbers
    pattern = r'(?<=\d)\xB7(?=\d)'

    for run in runs:
        original_text = run.text

        # Replace '•' with '.' only when between digits
        updated_text = re.sub(pattern, '.', original_text)

        if original_text != updated_text:
            run.text = updated_text
            changes.append((original_text, updated_text))

    for original, updated in changes:
        global_logs.append(
            f"[convert_decimal_to_baseline] Line {line_number}: '{original}' -> '{updated}'"
        )

    return runs



# change later
# Function to convert numbers to words (1 to 10)
def number_to_word(num):
    num_dict = {
        1: 'One', 2: 'Two', 3: 'Three', 4: 'Four', 5: 'Five',
        6: 'Six', 7: 'Seven', 8: 'Eight', 9: 'Nine', 10: 'Ten'
    }
    return num_dict.get(num, str(num))

# change later
# Function to convert words to numbers
def word_to_number(word):
    word_dict = {
        'one': 1, 'two': 2, 'three': 3, 'four': 4, 'five': 5,
        'six':6, 'seven':7, 'eight':8, 'nine':9, 'ten':10 
    }
    return word_dict.get(word.lower(), word)

# change later
# Function to process text and replace words with numbers, and numbers with words
def convert_text(text):
    
    text = re.sub(r'\b([1-9]|10)\b', lambda match: number_to_word(int(match.group(0))), text)
    text = re.sub(r'\b(one|two|three|four|five|six|seven|eight|nine|ten)\s*(kg|m|cm|g|l)\b', 
    lambda match: str(word_to_number(match.group(1))) + ' ' + match.group(2), text, flags=re.IGNORECASE)
    return text



# def adjust_punctuation_style_using_paragraph_text(text, para_runs):
#     """
#     Analyze `text` to detect italicized or bold characters followed by punctuation
#     and ensure the punctuation inherits the appropriate style (italic or bold).
#     """
#     for i in range(len(para_runs) - 1):
#         current_run = para_runs[i]
#         next_run = para_runs[i + 1]

#         # Check if current run ends with italicized text
#         if current_run.text and current_run.italic:
#             last_char = current_run.text[-1]
#             if next_run.text and next_run.text[0] in ".,!?\"'()":
#                 next_run.italic = True
        
#         elif current_run.text and current_run.bold:
#             last_char = current_run.text[-1]
#             if next_run.text and next_run.text[0] in ".,!?\"”'()":
#                 next_run.bold = True
    
#     # Return updated text after style adjustments
#     return text



def adjust_punctuation_style_using_paragraph_text(para_runs):
    """
    Analyze `para_runs` to detect italicized or bold characters followed by punctuation
    and ensure the punctuation inherits the appropriate style (italic or bold).
    """
    for i in range(len(para_runs) - 1):
        current_run = para_runs[i]
        next_run = para_runs[i + 1]

        # Check if current run ends with italicized text
        if current_run.text and current_run.italic:
            last_char = current_run.text[-1]
            if next_run.text and next_run.text[0] in ".,!?\"'()":
                next_run.italic = True
        
        # Check if current run ends with bold text
        elif current_run.text and current_run.bold:
            last_char = current_run.text[-1]
            if next_run.text and next_run.text[0] in ".,!?\"”'()":
                next_run.bold = True

    # Return updated runs after style adjustments
    return para_runs




# Dictionary to convert word numbers to integer values
word_to_num = {
    'zero': 0, 'one': 1, 'two': 2, 'three': 3, 'four': 4, 'five': 5,
    'six': 6, 'seven': 7, 'eight': 8, 'nine': 9, 'ten': 10, 'eleven': 11,
    'twelve': 12, 'thirteen': 13, 'fourteen': 14, 'fifteen': 15, 'sixteen': 16,
    'seventeen': 17, 'eighteen': 18, 'nineteen': 19, 'twenty': 20, 'twenty-one': 21,
    'twenty-two': 22, 'twenty-three': 23, 'twenty-four': 24, 'twenty-five': 25,
    'twenty-six': 26, 'twenty-seven': 27, 'twenty-eight': 28, 'twenty-nine': 29,
    'thirty': 30, 'thirty-one': 31, 'thirty-two': 32, 'thirty-three': 33,
    'thirty-four': 34, 'thirty-five': 35, 'thirty-six': 36, 'thirty-seven': 37,
    'thirty-eight': 38, 'thirty-nine': 39, 'forty': 40, 'forty-one': 41,
    'forty-two': 42, 'forty-three': 43, 'forty-four': 44, 'forty-five': 45,
    'forty-six': 46, 'forty-seven': 47, 'forty-eight': 48, 'forty-nine': 49,
    'fifty': 50, 'fifty-one': 51, 'fifty-two': 52, 'fifty-three': 53,
    'fifty-four': 54, 'fifty-five': 55, 'fifty-six': 56, 'fifty-seven': 57,
    'fifty-eight': 58, 'fifty-nine': 59, 'sixty': 60, 'sixty-one': 61,
    'sixty-two': 62, 'sixty-three': 63, 'sixty-four': 64, 'sixty-five': 65,
    'sixty-six': 66, 'sixty-seven': 67, 'sixty-eight': 68, 'sixty-nine': 69,
    'seventy': 70, 'seventy-one': 71, 'seventy-two': 72, 'seventy-three': 73,
    'seventy-four': 74, 'seventy-five': 75, 'seventy-six': 76, 'seventy-seven': 77,
    'seventy-eight': 78, 'seventy-nine': 79, 'eighty': 80, 'eighty-one': 81,
    'eighty-two': 82, 'eighty-three': 83, 'eighty-four': 84, 'eighty-five': 85,
    'eighty-six': 86, 'eighty-seven': 87, 'eighty-eight': 88, 'eighty-nine': 89,
    'ninety': 90, 'ninety-one': 91, 'ninety-two': 92, 'ninety-three': 93,
    'ninety-four': 94, 'ninety-five': 95, 'ninety-six': 96, 'ninety-seven': 97,
    'ninety-eight': 98, 'ninety-nine': 99, 'hundred': 100
}

# Reverse dictionary to convert integer values back to words
num_to_word = {v: k for k, v in word_to_num.items()}

# Function to convert word number to integer
def word_to_int(word):
    return word_to_num.get(word.lower(), None)

# Function to convert integer to word
def int_to_word(num):
    return num_to_word.get(num, None)

# Regular expression to match "word and word" pattern
pattern = re.compile(r'(\b\w+\b) and (\b\w+\b)')

# Function to process the string with regex and apply transformations
def process_string(text):
    def replace_match(match):
        word1 = match.group(1)
        word2 = match.group(2)
        
        # Convert words to their numeric values
        num1 = word_to_int(word1)
        num2 = word_to_int(word2)
        
        # If both numbers are less than 9, return them as word form
        if (num1 is not None and num1 < 9) and (num2 is not None and num2 < 9):
            return f"{word1} and {word2}"  # No change if both are < 9
        
        # If either number is greater than or equal to 9, convert to numeric form
        if (num1 is not None and num1 >= 9) or (num2 is not None and num2 >= 9):
            # Convert both to numeric form
            num1 = num1 if num1 is not None else word1
            num2 = num2 if num2 is not None else word2
            return f"{num1} and {num2}"  # Replace with numeric values
        
        return match.group(0)  # Return the match as is if both are < 9
    
    # Apply regex substitution with the replace function
    return pattern.sub(replace_match, text)
 


def highlight_and_correct(doc):
    chapter_counter = [0]
    line_number = 1
    abbreviation_dict = fetch_abbreviation_mappings()
    for para in doc.paragraphs:
        
        # replace_curly_quotes_with_straight(para.runs)

        # if para.text.strip().startswith("Chapter"):
        #     para.text = correct_chapter_numbering(para.text, chapter_counter)
        #     formatted_title = format_chapter_title(para.text)
        #     para.text = formatted_title

        # process_symbols_mark(para.runs, line_number)
        remove_commas_from_numbers(para.runs, line_number)
        # remove_spaces_from_four_digit_numbers(para.runs, line_number)
        # set_latinisms_to_roman_in_runs(para.runs, line_number)
        # convert_decimal_to_baseline(para.runs, line_number)

        # rename_section(para.runs)
        # replace_ampersand(para.runs)
        # correct_scientific_unit_symbols(para.runs)
        # adjust_ratios(para.runs)
        # format_dates(para.runs, line_number)
        # spell_out_number_and_unit_with_rules(para.runs, line_number)
        # remove_space_between_degree_and_direction(para.runs, line_number)
        # enforce_lowercase_units(para.runs, line_number)
        # precede_decimal_with_zero(para.runs, line_number)
        # format_ellipses_in_series(para.runs) # not added in log and not working
        # correct_possessive_names(para.runs, line_number)
        # use_numerals_with_percent(para.runs)
        # remove_concluding_slashes_from_urls(para.runs, line_number)
        # clean_web_addresses(para.runs)

        # apply_abbreviation_mapping(para.runs, abbreviation_dict, line_number)
        # apply_number_abbreviation_rule(para.runs, line_number)

        # format_titles_us_english_with_logging(para.runs)
        # units_with_bracket(para.runs)
        # correct_units_in_ranges_with_logging(para.runs, line_number)
        # correct_scientific_units_with_logging(para.runs)
        # replace_fold_phrases(para.runs)
        # correct_preposition_usage(para.runs)
        # correct_unit_spacing(para.runs)

        # remove_and(para.runs)
        # remove_quotation(para.runs)
        # convert_text(para.runs)

        # apply_quotation_punctuation_rule(para.runs)
        # enforce_dnase_rule(para.runs)

        # correct_acronyms(para.runs, line_number)
        # enforce_am_pm(para.runs, line_number)

        # enforce_eg_rule_with_logging(para.runs)
        # enforce_ie_rule_with_logging(para.runs)
        # enforce_serial_comma(para.runs)
        # apply_remove_italics_see_rule(para.runs)
        # process_string(para.runs)

        # standardize_etc(para.runs)
        # process_url_add_http(para.runs)
        # process_url_remove_http(para.runs)

        # lines = para.text.split('\n')
        # updated_lines = []
        # for line in lines:
        #     corrected_line = convert_century(line, line_number)
        #     updated_lines.append(corrected_line)
        #     line_number += 1

        # para.text = '\n'.join(updated_lines)
        # formatted_runs = []        
        # for run in para.runs:
        #     run_text = insert_thin_space_between_number_and_unit(run.text, line_number)

        #     words = run_text.split()
        #     for i, word in enumerate(words):
        #         original_word = word
        #         punctuation = ""

        #         if word[-1] in ",.?!:;\"'()[]{}":
        #             punctuation = word[-1]
        #             word = word[:-1]

        #         if (word.startswith('"') and word.endswith('"')) or (word.startswith("'") and word.endswith("'")):
        #             formatted_runs.append((original_word, None))
        #             if i < len(words) - 1:
        #                 formatted_runs.append((" ", None))
        #             continue

        #         if not word.strip():
        #             formatted_runs.append((original_word, None))
        #             if i < len(words) - 1:
        #                 formatted_runs.append((" ", None))
        #             continue

        #         if not us_dict.check(word.lower()):
        #             formatted_runs.append((original_word, RGBColor(255, 0, 0)))
        #         else:
        #             formatted_runs.append((original_word, None))

        #         if i < len(words) - 1:
        #             formatted_runs.append((" ", None))
        # para.clear()

        # for text, color in formatted_runs:
        #     adjusted_text = replace_straight_quotes_with_curly(text)
        #     new_run = para.add_run(adjusted_text)
        #     if color:
        #         new_run.font.color.rgb = color



def clean_word1(word):
    return ''.join(filter(str.isalnum, word)).lower()





# Helper function to extract text from docx file
def extract_text_from_docx(file_path):
    try:
        with open(file_path, "rb") as docx_file:
            result = mammoth.extract_raw_text(docx_file)
            return result.value
    except Exception as e:
        # logging.error(f"Error extracting text from file: {e}")
        return ""



@router.get("/process_us")
async def process_file(doc_id: int = Query(...)):
    try:
        # Connect to the database
        conn = get_db_connection()
        if conn is None:
            raise HTTPException(status_code=500, detail="Database connection error")
        
        cursor = conn.cursor()
        cursor.execute("SELECT * FROM row_document WHERE row_doc_id = %s", (doc_id,))
        rows = cursor.fetchone()

        if not rows:
            raise HTTPException(status_code=404, detail="Document not found")
        
        file_path = os.path.join(os.getcwd(), 'files', rows[1])

        # Verify the file exists
        if not os.path.exists(file_path):
            raise HTTPException(status_code=404, detail="File not found on server")

        # Start time of processing
        start_time = datetime.now()

        # Extract raw text using Mammoth (you need your own extract_text_from_docx method)
        file_content = extract_text_from_docx(file_path)
        text = file_content

        # Prepare log data for spell checking
        global_logs.append(f"FileName: {rows[1]}\n\n")

        # Split text into lines and process each line for spelling errors
        lines = text.split('\n')
        for index, line in enumerate(lines):
            words = line.split()
            for word in words:
                cleaned = clean_word1(word)
                if cleaned and not us_dict.check(cleaned):
                    suggestions = us_dict.suggest(cleaned)
                    suggestion_text = (
                        f" Suggestions: {', '.join(suggestions)}"
                        if suggestions else " No suggestions available"
                    )
                    global_logs.append(f"Line {index}: {word} ->{suggestion_text}")

        # End time and time taken
        end_time = datetime.now()
        time_taken = round((end_time - start_time).total_seconds(), 2)
        time_log = f"\nStart Time: {start_time}\nEnd Time: {end_time}\nAnalysis completed in {time_taken} seconds.\n\n"

        # Prepend the time log to the existing log data
        global_logs.insert(0, time_log)

        # Define the log filename based on the document ID and name
        document_name = rows[1].replace('.docx', '')
        log_filename = f"log_main.txt"
        
        # Define output path for the log file inside a directory based on doc_id
        output_path_file = Path(os.getcwd()) / 'output' / str(doc_id) / log_filename
        dir_path = output_path_file.parent

        # Ensure the output directory exists
        dir_path.mkdir(parents=True, exist_ok=True)

        # try:
        #     # Read existing content of the log file if exists
        #     if output_path_file.exists():
        #         with open(output_path_file, "r", encoding="utf-8") as log_file:
        #             existing_content = log_file.read()
        #         with open(output_path_file, "w", encoding="utf-8") as log_file:
        #             log_file.write(''.join(global_logs) + existing_content)
        #     else:
        #         # If the file doesn't exist, create it with the new log data
        #         with open(output_path_file, "w", encoding="utf-8") as log_file:
        #             log_file.write(''.join(global_logs))

        # except FileNotFoundError:
        #     # If the log file does not exist at all, create a new one
        #     with open(output_path_file, "w", encoding="utf-8") as log_file:
        #         log_file.write(''.join(global_logs))


        output_dir = os.path.join("output", str(doc_id))
        os.makedirs(output_dir, exist_ok=True)

        output_path = os.path.join(output_dir, f"processed_{os.path.basename(file_path)}")

        doc = docx.Document(file_path)
        highlight_and_correct(doc)
        doc.save(output_path)

        cursor.execute("SELECT final_doc_id FROM final_document WHERE row_doc_id = %s", (doc_id,))
        existing_rows = cursor.fetchall()

        if existing_rows:
            logging.info('File already processed in final_document. Skipping insert.')
        else:
            folder_url = f'/output/{doc_id}/'
            cursor.execute(
                '''INSERT INTO final_document (row_doc_id, user_id, final_doc_size, final_doc_url, status, creation_date)
                VALUES (%s, %s, %s, %s, %s, NOW())''',
                (doc_id, rows[1], rows[2], folder_url, rows[7])
            )
            logging.info('New file processed and inserted into final_document.')

        conn.commit()
        write_to_log(doc_id)
        logging.info(f"Processed file stored at: {output_path}")
        return {"success": True, "message": f"File processed and stored at {output_path}"}

    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

