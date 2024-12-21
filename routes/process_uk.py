import re
from docx.shared import RGBColor
from num2words import num2words
from word2number import w2n
import enchant
from fastapi import APIRouter, HTTPException, Depends, Query
from fastapi.responses import FileResponse
import os
import docx
from sqlalchemy import text
from db_config import get_db_connection
import mammoth
from datetime import datetime
from pathlib import Path
import logging  
import roman  # For converting Roman numerals to Arabic numerals
from roman import fromRoman


router = APIRouter()

# uk_dict = enchant.Dict("en_US") #for US english
uk_dict = enchant.Dict("en_GB") # for UK english

century_map = {
    1: "first",
    2: "second",
    3: "third",
    4: "fourth",
    5: "fifth",
    6: "sixth",
    7: "seventh",
    8: "eighth",
    9: "ninth",
    10: "tenth",
    11: "eleventh",
    12: "twelfth",
    13: "thirteenth",
    14: "fourteenth",
    15: "fifteenth",
    16: "sixteenth",
    17: "seventeenth",
    18: "eighteenth",
    19: "nineteenth",
    20: "twentieth",
    21: "twenty-first",
    22: "twenty-second",
    23: "twenty-third",
    24: "twenty-fourth",
    25: "twenty-fifth",
}



def replace_percent_with_symbol(text):
    """
    Replaces 'percent' or 'per cent' with '%' if preceded by a number.

    :param text: The text to process.
    :return: The modified text.
    """
    return re.sub(r"(\d+)\s?(percent|per cent)", r"\1%", text, flags=re.IGNORECASE)

def convert_century(word):
    match = re.match(r"(\d+)(st|nd|rd|th)$", word)
    if match:
        num = int(match.group(1))
        if num in century_map:
            return f"the {century_map[num]} century"
    return word

def clean_word(word):
    return word.strip(",.?!:;\"'()[]{}")

def replace_curly_quotes_with_straight(text):
    return text.replace("“", '"').replace("”", '"').replace("‘", "'").replace("’", "'")



# def replace_straight_single_quotes_with_curly(text):
#     text = re.sub(r"(^|\s)'", "‘", text)
#     text = re.sub(r"'(\s|$)", "’", text)
#     return text


def correct_acronyms(word):
    if re.match(r"([a-z]\.){2,}[a-z]\.?", word):
        return word.replace(".", "")
    if re.match(r"([A-Z]\.){2,}[A-Z]\.?", word):
        return word.replace(".", "")
    return word

def enforce_am_pm(word):
    word_lower = word.lower()
    if word_lower in {"am", "a.m", "pm", "p.m"}:
        if "a" in word_lower:
            return "a.m."
        elif "p" in word_lower:
            return "p.m."
    return word

def remove_unnecessary_apostrophes(word):
    word = re.sub(r"(\d{4})'s\b", r"\1s", word)
    word = re.sub(r"'(\d{2})s\b", r"\1s", word)
    word = re.sub(r"(\d{4}s)'\b", r"\1", word)
    word = re.sub(r"(\d+)'(s|st|nd|rd|th)\b", r"\1\2", word)
    word = re.sub(r"^(\d{2})s\b", r"19\1s", word)
    return word

def spell_out_number_and_unit(sentence):
    match = re.match(r"^(\d+)\s+([a-zA-Z]+)", sentence)
    if match:
        number = int(match.group(1))
        unit = match.group(2)
        number_word = num2words(number, to="cardinal")
        unit_word = unit.lower() if unit.lower()[-1] == 's' else unit.lower() + "s"
        return f"{number_word.capitalize()} {unit_word}{sentence[len(match.group(0)):]}"
    return sentence



def use_numerals_with_percent(text):
    """
    Converts spelled-out numbers with 'percent' or 'per cent' into numerals followed by '%'.
    """
    # Convert spelled-out numbers (like "five percent") to numerals with '%'
    text = re.sub(
        r"\b(one|two|three|four|five|six|seven|eight|nine|ten|eleven|twelve|thirteen|fourteen|fifteen|sixteen|seventeen|eighteen|nineteen|twenty|thirty|forty|fifty|sixty|seventy|eighty|ninety|hundred)\s?(percent|per cent)",
        lambda m: f"{w2n.word_to_num(m.group(1))}%",
        text,
        flags=re.IGNORECASE,
    )
    # Replace numeral with 'percent' or 'per cent' to numerals with '%'
    text = re.sub(r"(\d+)\s?(percent|per cent)", r"\1%", text, flags=re.IGNORECASE)
    return text



def correct_chapter_numbering(text, chapter_counter):
    """
    Ensures that chapter headings are numbered sequentially using an external counter.

    :param text: The text to process.
    :param chapter_counter: A list containing the chapter counter as its first element.
    :return: The modified text.
    """
    # Pattern to match chapter headings
    chapter_pattern = re.compile(r"(?i)\bchapter\s+((?:[IVXLCDM]+)|(?:[a-z]+)|(?:\d+))[:.]?\s")

    def replace_chapter_heading(match):
        """
        Replace matched chapter heading with sequential numbering.
        """
        chapter_counter[0] += 1  # Increment the shared counter
        return f"Chapter {chapter_counter[0]}: "

    # Apply the substitution across the text
    return chapter_pattern.sub(replace_chapter_heading, text)




# uncommet it later for point 
# def correct_scientific_unit_symbols(text):
#     """
#     Removes incorrect plural forms, apostrophes, or periods from scientific unit symbols.
#     Ensures unit symbols like 'kg', 'm', 'L', etc., are used properly.

#     :param text: The text to process.
#     :return: The modified text.
#     """
#     # List of common scientific unit symbols
#     units = [
#         "kg", "g", "mg", "L", "ml", "m", "cm", "mm", "km", "s", "min", "h", "A", 
#         "mol", "cd", "K", "Pa", "N", "J", "W", "C", "V", "Ω", "Hz", "Bq", "Gy", "Sv", "lx", "lm"
#     ]
#     # Create regex pattern for units followed by invalid characters
#     pattern = r"\b(\d+)\s?(" + "|".join(units) + r")['s.]?\b"
#     # Replace invalid forms with correct form
#     return re.sub(pattern, r"\1 \2", text)







def insert_thin_space_between_number_and_unit(text):
    thin_space = '\u2009'
    text = re.sub(r"(\d)(?=\s?[a-zA-Z]+)(?!\s?°)", r"\1" + thin_space, text)
    return text



def format_dates(text):
    text = re.sub(r"\b(\d+)\s?(BCE|CE)\b", lambda m: f"{m.group(1)} {m.group(2).lower()}", text)
    text = re.sub(r"\b(AD|BC)\.\b", r"\1 ", text)
    text = re.sub(r"(\d+)\s?(BCE|CE|AD|BC)\b", r"\1 \2", text)
    return text


def remove_space_between_degree_and_direction(text):
    return re.sub(r"(\d+) °\s?(N|S|E|W)\b", r"\1°\2", text)


def enforce_lowercase_k_and_l(text):
    text = re.sub(r"(\d) K([m|g])", r"\1 k\2", text)
    text = re.sub(r"\bL\b", "l", text)
    return text


def precede_decimal_with_zero(text):
    # Matches standalone decimal numbers below 1, e.g., ".75" or " .5", but not parts of larger numbers.
    text = re.sub(r"(?<!\d)(?<!\d\.)\.(\d+)", r"0.\1", text)
    return text


def adjust_terminal_punctuation_in_quotes(text):
    # Matches quoted matter ending with a question or exclamation mark, ensuring punctuation is inside the quotes.
    text = re.sub(
        r"([‘“])([^’”]*[?!])([’”])\.",
        r"\1\2\3",
        text
    )
    return text


def correct_possessive_names(text):
    # Handles singular possessives for names ending in 's'
    text = re.sub(r"\b([A-Za-z]+s)\b(?<!\bs')'", r"\1's", text)  # Add 's' for singular possessives
    text = re.sub(r"\b([A-Za-z]+s)'\b", r"\1'", text)  # Retain just the apostrophe for plurals
    return text


def remove_concluding_slashes_from_urls(text):
    # Matches URLs ending with a forward slash, but not when followed by other characters (like periods).
    text = re.sub(r"(https?://[^\s/]+(?:/[^\s/]+)*)/", r"\1", text)
    return text


def clean_web_addresses(text):
    return re.sub(r"<(https?://[^\s<>]+)>", r"\1", text)


def format_ellipses_in_series(text):
    # Matches series like "x1, x2, ..., xn" and ensures the ellipsis has a comma and space after it.
    text = re.sub(r"(\w+),\s*(\w+),\s*\.\.\.\s*(\w+)", r"\1, \2, …, \3", text)
    return text



def format_chapter_title(text):
    match = re.match(r"Chapter\s+([\dIVXLCDM]+)[\.:]\s*(.*)", text, re.IGNORECASE)
    if match:
        chapter_number = match.group(1)
        chapter_title = match.group(2).rstrip('.')  # Remove trailing period
        words = chapter_title.split()
        formatted_title = " ".join([
            word.capitalize() if i == 0 or len(word) >= 5 else word.lower()
            for i, word in enumerate(words)
        ])
        return f"{chapter_number}. {formatted_title}"
    return text



def uk_english_titles(text):
    """
    Formats titles for UK English, ensuring titles do not have dots
    and are correctly capitalized.
    """
    titles = {
        "mister": "Mr",
        "doctor": "Dr",
        "miss": "Miss",
        "mrs": "Mrs",
        "ms": "Ms",
        "professor": "Professor",
    }

    def replace_title(match):
        return titles.get(match.group(1).lower(), match.group(1))

    # Regular expression to match titles at word boundaries
    pattern = r"\b(" + "|".join(titles.keys()) + r")\b"
    text = re.sub(pattern, replace_title, text, flags=re.IGNORECASE)

    return text


import os
import re

def uk_english_titles_with_logging(text, doc_id):
    """
    Formats titles for UK English, ensuring titles do not have dots
    and are correctly capitalized, and logs changes to a file.

    :param text: The text to process.
    :param doc_id: The document ID for logging purposes.
    :return: The modified text.
    """
    titles = {
        "mister": "Mr",
        "doctor": "Dr",
        "miss": "Miss",
        "mrs": "Mrs",
        "ms": "Ms",
        "professor": "Professor",
    }

    # Create output directory if not exists
    output_dir = f'output/{doc_id}'
    os.makedirs(output_dir, exist_ok=True)

    # Prepare log file path
    log_file_path = os.path.join(output_dir, 'log_uk.txt')
    
    # List to store changes
    changes = []

    def replace_title(match):
        original_title = match.group(1)
        formatted_title = titles.get(original_title.lower(), original_title)
        if original_title != formatted_title:
            changes.append(f"Line {match.start()}: {original_title} -> {formatted_title}")
        return formatted_title

    # Regular expression to match titles at word boundaries
    pattern = r"\b(" + "|".join(titles.keys()) + r")\b"
    updated_text = re.sub(pattern, replace_title, text, flags=re.IGNORECASE)

    # Write changes to log file if any changes were made
    if changes:
        with open(log_file_path, 'w', encoding='utf-8') as log_file:
            log_file.write("\n".join(changes))
        print(f"Changes written to {log_file_path}")
    else:
        print("No changes detected.")

    # Return the updated text
    return updated_text







def highlight_and_correct(doc,doc_id):
    chapter_counter = [0] 
    for para in doc.paragraphs:
        if para.text.strip().startswith("Chapter"):
            # Update the paragraph text with sequential chapter numbering
            para.text = correct_chapter_numbering(para.text, chapter_counter)

            # Format the chapter title (optional)
            formatted_title = format_chapter_title(para.text)
            para.text = formatted_title
        
        
        para.text = format_dates(para.text)
        para.text = spell_out_number_and_unit(para.text)
        para.text = remove_space_between_degree_and_direction(para.text)
        para.text = enforce_lowercase_k_and_l(para.text)
        para.text = precede_decimal_with_zero(para.text)
        para.text = format_ellipses_in_series(para.text)  # New rule added here
        # para.text = adjust_terminal_punctuation_in_quotes(para.text)
        para.text = correct_possessive_names(para.text)
        # para.text = correct_scientific_unit_symbols(para.text)
        para.text = use_numerals_with_percent(para.text)
        para.text = replace_percent_with_symbol(para.text)
        para.text = remove_concluding_slashes_from_urls(para.text)
        para.text = clean_web_addresses(para.text)
        
        para.text = uk_english_titles_with_logging(para.text,doc_id)

        formatted_runs = []

        for run in para.runs:
            run_text = replace_curly_quotes_with_straight(run.text)
            run_text = insert_thin_space_between_number_and_unit(run_text)
            words = run_text.split()

            for i, word in enumerate(words):
                original_word = word
                punctuation = ""

                if word[-1] in ",.?!:;\"'()[]{}":
                    punctuation = word[-1]
                    word = word[:-1]

                if word.startswith('"') or word.startswith("'") or word.endswith('"') or word.endswith("'"):
                    formatted_runs.append((word, None))
                    if i < len(words) - 1:
                        formatted_runs.append((" ", None))
                    continue

                word = remove_unnecessary_apostrophes(word)

                cleaned_word = clean_word(word)
                corrected_word = cleaned_word

                if cleaned_word:
                    corrected_word = correct_acronyms(cleaned_word)
                    corrected_word = enforce_am_pm(corrected_word)
                    corrected_word = convert_century(corrected_word)

                    if corrected_word != cleaned_word:
                        formatted_runs.append((corrected_word + punctuation, RGBColor(0, 0, 0)))
                    elif not uk_dict.check(corrected_word.lower()):
                        formatted_runs.append((corrected_word + punctuation, RGBColor(255, 0, 0)))
                    else:
                        formatted_runs.append((corrected_word + punctuation, None))
                else:
                    formatted_runs.append((original_word + punctuation, None))

                if i < len(words) - 1:
                    formatted_runs.append((" ", None))

        para.clear()
        for text, color in formatted_runs:
            new_run = para.add_run(text)
            if color:
                new_run.font.color.rgb = color


def clean_word1(word):
    return ''.join(filter(str.isalnum, word)).lower()





# Helper function to extract text from docx file
def extract_text_from_docx(file_path):
    try:
        with open(file_path, "rb") as docx_file:
            result = mammoth.extract_raw_text(docx_file)
            return result.value
    except Exception as e:
        # logging.error(f"Error extracting text from file: {e}")
        return ""




@router.get("/process_uk")
async def process_file(doc_id: int = Query(...)):
    try:
        # Connect to the database
        conn = get_db_connection()
        if conn is None:
            raise HTTPException(status_code=500, detail="Database connection error")
        
        cursor = conn.cursor()
        cursor.execute("SELECT * FROM row_document WHERE row_doc_id = %s", (doc_id,))
        rows = cursor.fetchone()

        if not rows:
            raise HTTPException(status_code=404, detail="Document not found")
        
        # File path based on document ID
        file_path = os.path.join(os.getcwd(), 'files', rows[1])

        # Verify the file exists
        if not os.path.exists(file_path):
            raise HTTPException(status_code=404, detail="File not found on server")
        
        # Start time of processing
        start_time = datetime.now()

        # Extract raw text using Mammoth (you need your own extract_text_from_docx method)
        file_content = extract_text_from_docx(file_path)
        text = file_content

        # Prepare log data for spell checking
        log_data = []
        log_data.append(f"FileName: {rows[1]}\n\n")
        
        # Split text into lines and process each line for spelling errors
        lines = text.split('\n')
        for index, line in enumerate(lines):
            words = line.split()
            for word in words:
                cleaned = clean_word1(word)
                if cleaned and not uk_dict.check(cleaned):
                    suggestions = uk_dict.suggest(cleaned)
                    suggestion_text = (
                        f" Suggestions: {', '.join(suggestions)}"
                        if suggestions else " No suggestions available"
                    )
                    log_data.append(f"Line {index + 1}: {word} ->{suggestion_text}\n")

        # End time and time taken
        end_time = datetime.now()
        time_taken = round((end_time - start_time).total_seconds(), 2)
        time_log = f"\nStart Time: {start_time}\nEnd Time: {end_time}\nAnalysis completed in {time_taken} seconds.\n\n"

        # Define the log filename based on the document ID and name
        document_name = rows[1].replace('.docx', '')
        log_filename = f"{document_name}_log_us.txt"
        
        # Define output path for the log file inside a directory based on doc_id
        output_path_file = Path(os.getcwd()) / 'output' / str(doc_id) / log_filename
        dir_path = output_path_file.parent

        # Ensure the output directory exists
        dir_path.mkdir(parents=True, exist_ok=True)

        # Prepend the time log to the existing log data
        try:
            # Read existing content of the log file if exists
            if output_path_file.exists():
                with open(output_path_file, "r", encoding="utf-8") as log_file:
                    existing_content = log_file.read()
                with open(output_path_file, "w", encoding="utf-8") as log_file:
                    log_file.write(time_log + ''.join(log_data) + existing_content)
            else:
                # If the file doesn't exist, create it with the new log data
                with open(output_path_file, "w", encoding="utf-8") as log_file:
                    log_file.write(time_log + ''.join(log_data))

        except FileNotFoundError:
            # If the log file does not exist at all, create a new one
            with open(output_path_file, "w", encoding="utf-8") as log_file:
                log_file.write(time_log + ''.join(log_data))

        # Process the document and save it with corrections
        output_dir = os.path.join("output", str(doc_id))  # Folder based on doc_id
        os.makedirs(output_dir, exist_ok=True)

        output_path = os.path.join(output_dir, f"processed_{os.path.basename(file_path)}")

        # Process the document (assuming highlight_and_correct is your correction function)
        doc = docx.Document(file_path)
        highlight_and_correct(doc,doc_id)
        doc.save(output_path)

        # Save document metadata to the database if not already processed
        cursor.execute("SELECT final_doc_id FROM final_document WHERE row_doc_id = %s", (doc_id,))
        existing_rows = cursor.fetchall()

        if existing_rows:
            logging.info('File already processed in final_document. Skipping insert.')
        else:
            # Insert new record into final_document table
            folder_url = f'/output/{doc_id}/'
            cursor.execute(
                '''INSERT INTO final_document (row_doc_id, user_id, final_doc_size, final_doc_url, status, creation_date)
                VALUES (%s, %s, %s, %s, %s, NOW())''',
                (doc_id, rows[1], rows[2], folder_url, rows[7])
            )
            logging.info('New file processed and inserted into final_document.')

        # Commit changes to the database
        conn.commit()

        # Log the success message and return the response
        logging.info(f"Processed file stored at: {output_path}")
        return {"success": True, "message": f"File processed and stored at {output_path}"}

    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

